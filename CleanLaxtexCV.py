#create a python script to clean the latex cv

# import the necessary packages
import argparse
import math
import sys
import re
import shutil
import subprocess
import time
import glob
import csv
import random
import string
import datetime
import os 
import datetime
from docx import Document 


def process_student_thesis_titles(text,dossier_file_path): 
    # Extract master's thesis titles
    masters_data = extract_student_titles(dossier_file_path)
 
    # Extract PhD dissertation titles
    phd_data = extract_phd_titles(dossier_file_path)
  
    # Extract postdoc titles
    postdoc_data = extract_postdoc_titles(dossier_file_path)
    
    # Extract undergrad thesis titles
    undergrad_data = extract_undergrad_student_titles(dossier_file_path)
     
    # Merge the dictionaries
    student_data = {**masters_data, **phd_data, **postdoc_data, **undergrad_data}

    # Replace problematic characters in titles
    student_data = replace_problematic_characters_in_titles(student_data)
        
    # Update the LaTeX content with thesis titles - does not work for postdoc titles see next function
    text = add_title_to_name2(text, student_data)
       
    # reformat phd section for name, title, date
    text = reformat_phd_section(text)
    
    # reformat masters section for name, title, date    
    text = update_masters_section(text)
    
    # reformat postdoc section for name, title, date
    text = update_postdoc_section(text)
    
    # reformat undergrad section for name, title, date
    text = update_undergraduate_section(text)
        
    return text

def replace_problematic_characters_in_titles(student_data):
    problematic_char = u'\u2013'  # En dash (U+2013)

    for name, title in student_data.items():
        if problematic_char in title:
            #print(f"Found character '{problematic_char}' in {name}: {title}")
            student_data[name] = title.replace(problematic_char, '--')

    return student_data

def update_undergraduate_section(text_data):
    # Define section label
    label = "Undergraduate Honors\nThesis"
    
    # Extract the Undergraduate Honors Thesis section
    section_pattern = rf"\\subsubsection{{{label}}}.*?\\begin{{enumerate}}.*?\\end{{enumerate}}"
    section_match = re.search(section_pattern, text_data, re.DOTALL)
    
    if not section_match:
        print("Undergraduate Honors Thesis section not found!")
        return text_data
    
    section_content = section_match.group(0)
    
    # Extract individual items
    item_pattern = r"Undergraduate Honors Thesis\. \((.*?)\).\\\\\s+Advised: (.*?), \"(.*?)\","
    items = re.findall(item_pattern, section_content, re.DOTALL)
    
    # Format items
    formatted_items = []
    for time_period, name, title in items:
        formatted_item = f"\\item {name}, \"{title}\", {time_period}"
        formatted_items.append(formatted_item)
    
    # Combine formatted items
    formatted_section = "\\subsubsection{Undergraduate Honors\nThesis}\n\\begin{enumerate}\n\\def\\labelenumi{\\arabic{enumi}.}\n" + "\n".join(formatted_items) + "\n\\end{enumerate}"
    
    # Replace the old section with the new formatted section in the text data
    updated_text_data = text_data.replace(section_content, formatted_section)
    
    return updated_text_data

def update_postdoc_section(text_data):
    # Define section label
    label = "Postdoctoral Mentorship"
    
    # Extract the Postdoctoral Mentorship section
    section_pattern = rf"\\subsubsection{{{label}}}.*?\\begin{{enumerate}}.*?\\end{{enumerate}}"
    section_match = re.search(section_pattern, text_data, re.DOTALL)
    
    if not section_match:
        print("Postdoctoral Mentorship section not found!")
        return text_data
    
    section_content = section_match.group(0)
    
    # Extract individual items
    item_pattern = r"Postdoctoral Mentorship\. \((.*?)\).\\\\\s+Advised: (.*?), \"(.*?)\","
    items = re.findall(item_pattern, section_content, re.DOTALL)
    
    # Format items
    formatted_items = []
    for time_period, name, title in items:
        formatted_item = f"\\item {name}, \"{title}\", {time_period}"
        formatted_items.append(formatted_item)
    
    # Combine formatted items
    formatted_section = "\\subsubsection{Postdoctoral Mentorship}\n\\begin{enumerate}\n\\def\\labelenumi{\\arabic{enumi}.}\n" + "\n".join(formatted_items) + "\n\\end{enumerate}"
    
    # Replace the old section with the new formatted section in the text data
    updated_text_data = text_data.replace(section_content, formatted_section)
    
    return updated_text_data

def update_masters_section(text_data):
    # Define the two possible formats for "Master's Thesis"
    master_thesis_options = ["Master's Thesis", "Master\\textquotesingle s Thesis"]

    # Detect which format is used in the text
    master_thesis_format = next((option for option in master_thesis_options if option in text_data), None)

    if master_thesis_format is None:
        print("Master's Thesis section format not detected!")
        return text_data

    # Escape backslashes for regex
    escaped_master_thesis_format = master_thesis_format.replace('\\', '\\\\')

    # Prepare regex pattern for the section
    section_pattern = rf"(\\subsubsection\{{.*?{escaped_master_thesis_format}.*?\}}.*?\\begin\{{enumerate\}})(.*?)(\\end\{{enumerate\}})"
    
    section_match = re.search(section_pattern, text_data, re.DOTALL)

    if not section_match:
        print("Master's Thesis section not found!")
        return text_data

    section_start, section_content, section_end = section_match.groups()
    
    # Adjusting item pattern for matching items within the section
    item_pattern = rf"\\item\s*{escaped_master_thesis_format}\.\s*\((.*?)\)\.\\\\\s*Advised:\s*(.*?),\s*\"(.*?)\""
    
    items = re.findall(item_pattern, section_content, re.DOTALL | re.MULTILINE)
    
    if not items:
        print("No items found in Master's Thesis section!")
        return text_data

    # Format items
    formatted_items = [f"\\item {name}, \"{title}\", {time_period}" for time_period, name, title in items]

    # Combine formatted items
    formatted_section_content = "\n".join(formatted_items)

    # Replace the old section with the new formatted section in the text data
    updated_section = section_start + formatted_section_content + section_end
    updated_text_data = text_data.replace(section_content, formatted_section_content)
    
    return updated_text_data

def reformat_masters_section(text_data):
    # Define section label
    label = "Master's Thesis"
    
    # Extract the Master's Thesis section
    section_pattern = rf"(\\subsubsection{{{label}.*?}}.*?\\begin{enumerate}.*?\\end{enumerate})"
    section_match = re.search(section_pattern, text_data, re.DOTALL)
    
    if not section_match:
        return text_data
    
    section_content = section_match.group(1)
    
    # Pattern to extract student details
    student_pattern = rf"{label}. \((.*?)\).+?Advised: (.*?), \"(.*?)\","
    student_matches = re.findall(student_pattern, section_content, re.DOTALL)
    
    # Create a new formatted section
    new_section = f"\\subsubsection{{{label}}}\n\n\\begin{enumerate}\n"
    for time_period, name, title in student_matches:
        # Split name into first and last names
        names = name.split()
        first_name = names[0]
        last_name = names[-1]
        
        # Add to the new section
        new_section += f"\\item\n  {first_name} {last_name}, \"{title}\", {time_period}\n"
    
    new_section += "\\end{enumerate}\n"
    
    # Replace the old section with the new one
    text_data = text_data.replace(section_content, new_section)
    
    return text_data

def reformat_sections(text_data):
    # Define section labels
    section_labels = {
        "postdoc": "Postdoctoral Mentorship",
        "undergrad": "Undergraduate Honors Thesis",
        "masters": "Master's Thesis"
    }
    
    # Function to reformat a specific section
    def reformat_section(section_content, label):
        student_pattern = rf"{label}. \((.*?)\).+?Advised: (.*?), \"(.*?)\","
        student_matches = re.findall(student_pattern, section_content, re.DOTALL)
        
        # Create a new formatted section
        new_section = f"\\subsubsection{{{label}}}\n\n\\begin{enumerate}\n"
        for time_period, name, title in student_matches:
            # Split name into first and last names
            names = name.split()
            first_name = names[0]
            last_name = names[-1]
            
            # Add to the new section
            new_section += f"\\item\n  {first_name} {last_name}, \"{title}\", {time_period}\n"
        
        new_section += "\\end{enumerate}\n"
        return new_section
    
    # Iterate over each section label and reformat
    for section, label in section_labels.items():
        section_pattern = rf"(\\subsubsection{{{label}.*?}}.*?\\begin{enumerate}.*?\\end{enumerate})"
        section_match = re.search(section_pattern, text_data, re.DOTALL)
        
        if section_match:
            section_content = section_match.group(1)
            new_section = reformat_section(section_content, label)
            text_data = text_data.replace(section_content, new_section)
    
    return text_data

def extract_undergrad_student_titles(file_path):
    #file_path = r"C:\Users\rhk12\OneDrive - The Pennsylvania State University\resume\CV\20231021-095357-CDT.docx"

    # Check if the file exists
    if not os.path.exists(file_path):
        return f"File not found at '{file_path}'"
    
    # If the file exists, proceed to read its content
    doc = Document(file_path)
    full_text = [para.text for para in doc.paragraphs]

    # Define the sub-section titles
    sub_section_title = "Undergraduate Honors Thesis Advisor"
    end_section_title = "THE SCHOLARSHIP OF Research and Creative Accomplishments"

    # Find the start of the sub-section
    start_index = next((i for i, text in enumerate(full_text) if sub_section_title in text), None)
    if start_index is None:
        return {}

    # Find the end of the sub-section
    end_index = next((i for i, text in enumerate(full_text[start_index:]) if end_section_title in text), None)
    if end_index is None:
        end_index = len(full_text)
    else:
        end_index += start_index

    # Extract undergrad student names and thesis titles
    undergrad_data = {}
    for line in full_text[start_index + 1:end_index]:
        if ", Undergraduate." in line:
            name, title_with_dates = line.split(", Undergraduate.", 1)
            title = title_with_dates.split(". (")[0].strip()
            
            # Remove the "Date Graduated" part from the title
            title = re.sub(r"\. Date Graduated:.*$", "", title)
            
            undergrad_data[name.strip()] = title

    return undergrad_data

def add_postdoc_work(text, postdoc_data):
    for name, title in postdoc_data.items():
        # Split the name into first and last names
        last_name, first_initial = name.split(', ')
        first_initial = first_initial.strip()[0]  # Get the first character of the first name
        
        # Construct two patterns: one for full name and one for first initial
        full_name_pattern = rf"(\\subsubsection\{{Postdoctoral Mentorship\}}.*?Advised: ){first_initial}.*?{last_name}"
        initial_pattern = rf"(\\subsubsection\{{Postdoctoral Mentorship\}}.*?Advised: ){first_initial} {last_name}"
        
        # Try matching with the full name
        if re.search(full_name_pattern, text, flags=re.DOTALL):
            text = re.sub(full_name_pattern, f"\\1{first_initial}. {last_name}, \"{title}\"", text, flags=re.DOTALL)
            print(f"Updated {first_initial}. {last_name}'s entry.")  # Debugging print statement
        # If that doesn't work, try matching with the first initial
        elif re.search(initial_pattern, text, flags=re.DOTALL):
            text = re.sub(initial_pattern, f"\\1{first_initial}. {last_name}, \"{title}\"", text, flags=re.DOTALL)
            print(f"Updated {first_initial}. {last_name}'s entry using initial.")  # Debugging print statement
        else:
            print(f"Pattern not found for: {first_initial}. {last_name} or {first_initial} {last_name}")  # Debugging print statement

    return text

def reformat_phd_section(text_data):
    # Extract the Ph.D. section
    phd_section_pattern = r"(\\subsubsection{Ph\.D\. Dissertation}.*?\\begin{enumerate}.*?\\end{enumerate})"
    phd_section_match = re.search(phd_section_pattern, text_data, re.DOTALL)
    
    if not phd_section_match:
        print("Ph.D. section not found in the provided text.")
        return text_data
    
    phd_section = phd_section_match.group(1)
    
    # Extract details for each student
    student_pattern = r"Ph\.D\. Dissertation\. \((.*?)\).+?Advised: (.*?), \"(.*?)\","
    student_matches = re.findall(student_pattern, phd_section, re.DOTALL)
    
    # Create a new formatted Ph.D. section
    new_phd_section = "\\subsubsection{Ph.D. Dissertation}\n\n\\begin{enumerate}\n"
    for time_period, name, title in student_matches:
        # Split name into first and last names
        names = name.split()
        first_name = names[0]
        rest_of_name = " ".join(names[1:])
        
        # Add to the new Ph.D. section
        new_phd_section += f"\\item\n  {first_name} {rest_of_name}, \"{title}\", {time_period}\n"

    
    new_phd_section += "\\end{enumerate}\n"
    
    # Replace the original Ph.D. section with the new one
    updated_text_data = text_data.replace(phd_section, new_phd_section)
    
    return updated_text_data

def add_title_to_name(text, data):
    for name, title in data.items():
        # Adjust for the name format in the postdoc section
        if ',' in name:
            last_name, first_initial = name.split(', ')
            formatted_name_full = f"{first_initial}. {last_name}"
            formatted_name_last = last_name
        else:
            formatted_name_full = name
            formatted_name_last = name.split()[0]  # Just the first name
        
        # Check which naming convention exists in the LaTeX document
        if formatted_name_full in text:
            pattern = re.escape(formatted_name_full)
        elif formatted_name_last in text:
            pattern = re.escape(formatted_name_last)
        else:
            print(f"Pattern not found for: {name}")
            continue
        
        # Extract the time period
        time_period_match = re.search(r"\((.*?)\)", text)
        if time_period_match:
            time_period = time_period_match.group(1)
        else:
            time_period = ""

        # If we're dealing with the postdoc section
        if "Postdoctoral Mentorship" in title:
            postdoc_section_pattern = r"(\\subsubsection\{Postdoctoral Mentorship\}.*?)" + pattern
            match = re.search(postdoc_section_pattern, text, re.DOTALL)
            if match:
                text = text.replace(match.group(0), match.group(0).replace(f"Advised: {formatted_name_full}", f"{formatted_name_last}, \"{title}\", {time_period}"))
        # If we're dealing with the undergraduate section
        elif "Undergraduate Honors Thesis" in title:
            undergrad_section_pattern = r"(\\subsubsection\{Undergraduate Honors\s*Thesis\}.*?Advised: )" + pattern
            match = re.search(undergrad_section_pattern, text, re.DOTALL)
            if match:
                text = text.replace(match.group(0), match.group(0).replace(f"Advised: {formatted_name_full}", f"{formatted_name_last}, \"{title}\", {time_period}"))
        # For Master's and Ph.D. students
        else:
            if re.search(pattern, text):
                text = re.sub(pattern, f"{formatted_name_last}, \"{title}\", {time_period}", text)
    return text

def add_title_to_name2(text, data):
    # Extract the section between \subsection{Directed Student Learning} and \subsection{Teaching Experience}
    start_index = text.find(r'\subsection{Directed Student Learning}')
    end_index = text.find(r'\subsection{Teaching Experience}', start_index)
    
    if start_index == -1 or end_index == -1:
        print("Directed Student Learning or Teaching Experience section not found!")
        return text

    section = text[start_index:end_index]

    for name, title in data.items():
        # Adjust for the name format in the postdoc section
        if ',' in name:
            last_name, first_initial = name.split(', ')
            formatted_name_full = f"{first_initial}. {last_name}"
            formatted_name_last = last_name
        else:
            formatted_name_full = name
            formatted_name_last = name.split()[0]  # Just the first name
        
        # Check which naming convention exists in the section
        if formatted_name_full in section:
            pattern = re.escape(formatted_name_full)
        elif formatted_name_last in section:
            pattern = re.escape(formatted_name_last)
        else:
            print(f"Pattern not found for: {name}")
            continue
        
        # Extract the time period
        time_period_match = re.search(r"\((.*?)\)", section)
        if time_period_match:
            time_period = time_period_match.group(1)
        else:
            time_period = ""

        # If we're dealing with the postdoc section
        if "Postdoctoral Mentorship" in title:
            postdoc_section_pattern = r"(\\subsubsection\{Postdoctoral Mentorship\}.*?)" + pattern
            match = re.search(postdoc_section_pattern, section, re.DOTALL)
            if match:
                section = section.replace(match.group(0), match.group(0).replace(f"Advised: {formatted_name_full}", f"{formatted_name_last}, \"{title}\", {time_period}"))
        # If we're dealing with the undergraduate section
        elif "Undergraduate Honors Thesis" in title:
            undergrad_section_pattern = r"(\\subsubsection\{Undergraduate Honors\s*Thesis\}.*?Advised: )" + pattern
            match = re.search(undergrad_section_pattern, section, re.DOTALL)
            if match:
                section = section.replace(match.group(0), match.group(0).replace(f"Advised: {formatted_name_full}", f"{formatted_name_last}, \"{title}\", {time_period}"))
        # For Master's and Ph.D. students
        else:
            if re.search(pattern, section):
                section = re.sub(pattern, f"{formatted_name_last}, \"{title}\", {time_period}", section)

    # Replace the original section with the modified one
    text = text[:start_index] + section + text[end_index:]
    
    return text

def add_undergrad_titles(text, undergrad_data):
    for name, title in undergrad_data.items():
        # Convert the name format from 'Last, F.' to 'First Last'
        last_name, first_initial = name.split(', ')
        # We'll use a placeholder for the first name since we don't have the full first name
        first_name_placeholder = f"{first_initial}.*"
        
        # Define the pattern to search for the student's entry in the LaTeX content
        pattern = rf"(Advised: {first_name_placeholder} {last_name})"
        
        # Check if the pattern exists in the text
        if re.search(pattern, text):
            # Replace the pattern with the pattern + title
            text =  re.sub(pattern, f"\\1, \"{title}\"", text)
            print(f"Updated {first_initial}. {last_name}'s entry.")
        else:
            print(f"Pattern for {first_initial}. {last_name} not found in the text!")
    
    return text

def extract_postdoc_titles(file_path):
    #file_path = r"C:\Users\rhk12\OneDrive - The Pennsylvania State University\resume\CV\20231021-095357-CDT.docx"
    
    # Check if the file exists
    if not os.path.exists(file_path):
        return f"File not found at '{file_path}'"
    
    # If the file exists, proceed to read its content
    doc = Document(file_path)
    full_text = [para.text for para in doc.paragraphs]

    # Define the sub-section titles
    sub_section_title = "Postdoctoral Mentorship Advisor"
    end_section_title = "Research Activity Advisor"

    # Find the start of the sub-section
    start_index = next((i for i, text in enumerate(full_text) if sub_section_title in text), None)
    if start_index is None:
        return {}

    # Find the end of the sub-section
    end_index = next((i for i, text in enumerate(full_text[start_index:]) if end_section_title in text), None)
    if end_index is None:
        end_index = len(full_text)
    else:
        end_index += start_index

    # Extract postdoc names and work titles
    postdoc_data = {}
    for line in full_text[start_index + 1:end_index]:
        if ". " in line:
            name, title_with_dates = line.split(". ", 1)
            title = title_with_dates.split(". (")[0].strip()
            postdoc_data[name.strip()] = title

    return postdoc_data

def extract_phd_titles(file_path):
    #file_path = r"C:\Users\rhk12\OneDrive - The Pennsylvania State University\resume\CV\20231021-095357-CDT.docx"
    
    # Check if the file exists
    if not os.path.exists(file_path):
        return f"File not found at '{file_path}'"
    
    # If the file exists, proceed to read its content
    doc = Document(file_path)
    full_text = [para.text for para in doc.paragraphs]

    # Define the sub-section titles
    sub_section_title = "Ph.D. Dissertation Advisor"
    end_section_title = "Ph.D. Dissertation Committee Member"  # This is the section that follows the Advisor section

    # Find the start of the sub-section
    start_index = next((i for i, text in enumerate(full_text) if sub_section_title in text), None)
    if start_index is None:
        print(f"Section '{sub_section_title}' not found")
        return {}

    # Find the end of the sub-section
    end_index = next((i for i, text in enumerate(full_text[start_index + 1:]) if end_section_title in text), None)
    if end_index is None:
        end_index = len(full_text)
    else:
        end_index += start_index + 1

    # Extract student names and dissertation titles
    student_data = {}
    for line in full_text[start_index + 1:end_index]:
        if ", " in line and "Ph.D." in line:
            name = line.split(",")[0].strip()
            title_start = line.find("Ph.D.") + 6
            title_end = line.find(".", title_start)
            title = line[title_start:title_end].strip()
            student_data[name] = title

    return student_data

def extract_student_titles(file_path):
    #file_path = r"C:\Users\rhk12\OneDrive - The Pennsylvania State University\resume\CV\20231021-095357-CDT.docx"
    
    # Check if the file exists
    if not os.path.exists(file_path):
        return f"File not found at '{file_path}'"
    
    # If the file exists, proceed to read its content
    doc = Document(file_path)
    full_text = [para.text for para in doc.paragraphs]

    # Define the sub-section titles
    sub_section_titles = ["Master's Thesis Advisor", "Master\u2019s Thesis Advisor"]
    end_section_titles = ["Master's Thesis Committee Member", "Master\u2019s Thesis Committee Member"]

    # Find the start of the sub-section
    start_index = None
    for title in sub_section_titles:
        start_index = next((i for i, text in enumerate(full_text) if title in text), None)
        if start_index is not None:
            break

    if start_index is None:
        print(f"Section 'Master's Thesis Advisor' not found")
        return

    # Find the end of the sub-section
    end_index = None
    for title in end_section_titles:
        end_index = next((i for i, text in enumerate(full_text[start_index:]) if title in text), None)
        if end_index is not None:
            break

    if end_index is None:
        end_index = len(full_text)
    else:
        end_index += start_index

    # Extract student names and thesis titles
    student_data = {}
    for line in full_text[start_index + 1:end_index]:
        if ", " in line and "MS." in line:
            name = line.split(",")[0].strip()
            title_start = line.find("MS.") + 4
            title_end = line.find(".", title_start)
            title = line[title_start:title_end].strip()
            student_data[name] = title

    return student_data

def extract_subsubsection(text, title):
    # Use regex to extract the subsubsection based on the title
    pattern = re.compile(r'\\subsubsection{' + re.escape(title) + r'}.*?\\end{enumerate}', re.DOTALL)
    match = pattern.search(text)
    if match:
        return match.group(0)
    else:
        print(f"Title '{title}' not found in extract_subsubsection!")
        return None

def reorder_student_sections(text):
    # Define the desired order
    order = ["Ph.D. Dissertation", "Master's Thesis", "Postdoctoral Mentorship", "Undergraduate Honors\nThesis"]
    
    # Extract each subsubsection based on the order
    sections = [extract_subsubsection(text, title) for title in order]
    
    # Check if any section is missing and remove it
    sections = [section for section in sections if section is not None]
    
    # Concatenate the reordered sections
    reordered_sections = "\n".join(sections)
    
    # Replace the original Directed Student Learning section with the reordered one
    start_idx = text.find(r'\subsection{Directed Student Learning}')
    end_idx = text.find(r'\subsection', start_idx + 1)
    if end_idx == -1:
        end_idx = len(text)
    
    return text[:start_idx] + r'\subsection{Directed Student Learning}' + '\n' + reordered_sections + text[end_idx:]
def reorder_student_sections4(text):
    # Define the two possible formats for "Master's Thesis"
    master_thesis_options = ["Master's Thesis", "Master\\textquotesingle s Thesis"]

    # Detect which format is used in the text
    master_thesis_format = next((option for option in master_thesis_options if option in text), None)

    if master_thesis_format is None:
        print("Master's Thesis section format not detected")
        return text

    # Define the desired order based on the detected format
    order = ["Ph.D. Dissertation", master_thesis_format, "Postdoctoral Mentorship", "Undergraduate Honors\nThesis"]

    # Extract each subsubsection based on the order
    sections = [extract_subsubsection(text, title) for title in order]

    # Check if any section is missing and remove it
    sections = [section for section in sections if section is not None]

    # Concatenate the reordered sections
    reordered_sections = "\n".join(sections)

    # Replace the original Directed Student Learning section with the reordered one
    start_idx = text.find(r'\subsection{Directed Student Learning}')
    end_idx = text.find(r'\subsection', start_idx + 1)
    if end_idx == -1:
        end_idx = len(text)

    return text[:start_idx] + r'\subsection{Directed Student Learning}' + '\n' + reordered_sections + text[end_idx:]
def reorder_student_sections3(text):
    # Define the two possible orders for "Master's Thesis"
    order1 = ["Ph.D. Dissertation", "Master's Thesis", "Postdoctoral Mentorship", "Undergraduate Honors\nThesis"]
    order2 = ["Ph.D. Dissertation", "Master\\textquotesingle s Thesis", "Postdoctoral Mentorship", "Undergraduate Honors\nThesis"]

    # Extract each subsubsection based on the two orders
    sections1 = [extract_subsubsection(text, title) for title in order1]
    sections2 = [extract_subsubsection(text, title) for title in order2]

    # Combine the two sets of sections, removing any None entries
    combined_sections = [section for section in sections1 + sections2 if section is not None]

    # Remove duplicate sections (if any)
    seen_titles = set()
    unique_sections = []
    for section in combined_sections:
        title = section.split('\n', 1)[0]  # Extract the title of the section
        if title not in seen_titles:
            seen_titles.add(title)
            unique_sections.append(section)

    # Concatenate the reordered sections
    reordered_sections = "\n".join(unique_sections)
    
    # Replace the original Directed Student Learning section with the reordered one
    start_idx = text.find(r'\subsection{Directed Student Learning}')
    end_idx = text.find(r'\subsection', start_idx + 1)
    if end_idx == -1:
        end_idx = len(text)
    
    return text[:start_idx] + r'\subsection{Directed Student Learning}' + '\n' + reordered_sections + text[end_idx:]
def reorder_student_sections2(text):
    # Define the desired order with both variations for "Master's Thesis"
    order = ["Ph.D. Dissertation", "Master's Thesis", "Master\\textquotesingle s Thesis", "Postdoctoral Mentorship", "Undergraduate Honors\nThesis"]
    
    # Extract each subsubsection based on the order
    sections = [extract_subsubsection(text, title) for title in order]
    
    # Check if any section is missing and remove it
    sections = [section for section in sections if section is not None]

    # Remove duplicate "Master's Thesis" sections (if both variations are present)
    master_thesis_sections = [section for section in sections if "Master's Thesis" in section or "Master\\textquotesingle s Thesis" in section]
    if len(master_thesis_sections) > 1:
        sections = [section for section in sections if section not in master_thesis_sections]
        sections.insert(1, master_thesis_sections[0])  # Keep only the first occurrence

    # Concatenate the reordered sections
    reordered_sections = "\n".join(sections)
    
    # Replace the original Directed Student Learning section with the reordered one
    start_idx = text.find(r'\subsection{Directed Student Learning}')
    end_idx = text.find(r'\subsection', start_idx + 1)
    if end_idx == -1:
        end_idx = len(text)
    
    return text[:start_idx] + r'\subsection{Directed Student Learning}' + '\n' + reordered_sections + text[end_idx:]


def boldface_name_in_publications(text):
    # Extract the publications section
    start_publications = text.find(r'\subsection{Publications}\label{publications}')
    start_presentations = text.find(r'\subsection{Presentations}\label{presentations}')
    
    if start_publications == -1 or start_presentations == -1:
        print("Publications or Presentations subsection not found!")
        return text

    publications_section = text[start_publications:start_presentations]
    
    # Boldface the name in the publications section
    patterns = [r'(Kraft, R\. H\.)', r'(Kraft,)', r'(Kraft, R\.)']
    for pattern in patterns:
        publications_section = re.sub(pattern, r'\\textbf{\1}', publications_section)
    
    # Replace the old publications section with the modified one
    modified_text = text[:start_publications] + publications_section + text[start_presentations:]
    
    return modified_text

def reorder_publications(content):
    # Extract the content from \subsection{Publications} up to \subsection{Presentations}
    pattern = r'(\\subsection\{Publications\}.*?)(\\subsection\{Presentations\})'
    match = re.search(pattern, content, re.DOTALL)
    
    if not match:
        return content  # No change if the pattern isn't found
    
    publications_section = match.group(1)
    presentations_label = match.group(2)
    
    # Split into individual subsubsections
    subsubsections = re.findall(r'(\\subsubsection\{.*?\}.*?)(?=\\subsubsection|$)', publications_section, re.DOTALL)
    
    # Define the desired order
    order = ["Journal Article", "Conference Proceeding", "Book Chapters", "Other"]
    
    # Create a dictionary with subsubsection titles as keys and their content as values
    subsubsection_dict = {}
    for subsubsection in subsubsections:
        title = re.search(r'\\subsubsection\{(.*?)\}', subsubsection).group(1)
        subsubsection_dict[title] = subsubsection

    # Construct the reordered content based on the desired order
    sorted_content = ''.join([subsubsection_dict[title] for title in order if title in subsubsection_dict])
    
    # Replace old subsubsections with sorted ones in the publications section
    updated_publications = publications_section.replace(''.join(subsubsections), sorted_content)
    
    # Construct the updated content
    return content.replace(publications_section, updated_publications)

    
def capitalize_subsections(text_content):
    # Find all \subsection{...} patterns, including those spanning multiple lines
    subsections = re.findall(r'\\subsection\{(.*?)\}', text_content, re.DOTALL)

    # For each found subsection, replace it with its capitalized version
    for subsection in subsections:
        # Remove newline and excess spaces
        subsection_cleaned = ' '.join(subsection.split())
        capitalized_subsection = subsection_cleaned.upper()
        # Replace the original with the cleaned and capitalized version
        text_content = re.sub(
            r'\\subsection\{' + re.escape(subsection) + r'\}',
            r'\\subsection{' + capitalized_subsection + '}',
            text_content,
            flags=re.DOTALL
        )

    return text_content

def add_date_to_header(text_content):
    # Get the current month and year
    now = datetime.datetime.now()
    month_year = now.strftime("%B %Y")

    # Define the header settings with a darker gray color
    header_settings = r"""
\usepackage{fancyhdr}
\usepackage{xcolor}
\definecolor{darkgray}{gray}{0.4} % Define a darker gray color
\fancypagestyle{firstpage}{
    \fancyhf{} % Clear all headers and footers first
    \rhead{\textcolor{darkgray}{""" + month_year + r"""}}
    \renewcommand{\headrulewidth}{0pt} % No header rule
}
\thispagestyle{firstpage}
"""

    # Insert the header settings right before \begin{document}
    marker = r"\begin{document}"
    preamble_end = text_content.find(marker)
    
    if preamble_end != -1:
        modified_content = text_content[:preamble_end] + header_settings + text_content[preamble_end:]
        return modified_content
    else:
        print("Warning: Marker for preamble end (\\begin{document}) not found in the text content.")
        return text_content  # Return the original content if marker is not found
def format_header(text):
    # Define the block of text to be removed
    to_remove = r"""
\textbf{Dr. Reuben H. Kraft}\\
The Pennsylvania State University\\
EN - Mechanical Engineering\\
(814) 867-4570\\
Email: rhk12@psu.edu
""".strip()  # .strip() is used to remove any leading/trailing whitespace

    # Replace the block with an empty string
    text = text.replace(to_remove, '')

    # Add the desired formatted name after \begin{document}
    new_header = r"""
\begin{center}
\LARGE \textbf{\textsc{REUBEN H. KRAFT}} \\
\rule{\linewidth}{2pt}
\end{center}
"""
    # Add a line space
    new_header += '\n\n'
    new_header += r"\normalsize % Return to the default font size"
    
    # Find the position of \begin{document} and insert the new header after it
    insert_pos = text.find(r'\begin{document}') + len(r'\begin{document}')
    text = text[:insert_pos] + '\n\n' + new_header + '\n\n' + text[insert_pos:]

    return text
def add_custom_package(text, package_name="mystyle"):
    # Find the position of \author{}
    insert_pos = text.find(r'\author{')
    
    # If \author{} is found, insert the \usepackage command before it
    if insert_pos != -1:
        text = text[:insert_pos] + f"\\usepackage{{{package_name}}}\n" + text[insert_pos:]
    
    return text
def set_section_colors(text):
    # Define the LaTeX commands to change the colors
    color_commands = r"""
\usepackage{titlesec}
\usepackage{color}
\titleformat{\subsection}
  {\normalfont\large\bfseries\color{blue}} % format
  {\thesubsection} % label
  {1em} % sep
  {} % before-code
\titleformat{\subsubsection}
  {\normalfont\normalsize\bfseries\color{red}} % format
  {\thesubsubsection} % label
  {1em} % sep
  {} % before-code
"""

    # Insert the color commands before \begin{document}
    insert_pos = text.find(r'\begin{document}')
    if insert_pos != -1:
        text = text[:insert_pos] + color_commands + text[insert_pos:]

    return text

def process_courses(text):
    course_descriptions = {
        "330": "Computational Tools for Engineers (ME 330)",
        "360": "Machine Design (ME 360)",
        "461": "Introduction to Finite Element Analysis (ME 461)",
        "563": "Nonlinear Finite Element Analysis (ME 563)",
        "497": "Development course for Computational Tools for Engineers (ME 497)",
        "440": "Capstone Design (ME 440)",
    }
    excluded_courses = {"600", "596", "496", "494", "610"}

    # Find the start and end of the course listings section
    start_of_section = text.find(r'\subsection{Teaching Experience}\label{teaching-experience}')
    end_of_section = text.find(r'\subsection{Service}\label{service}')

    # If the start or end markers are not found, return the original text
    if start_of_section == -1 or end_of_section == -1:
        print("Warning: Course section markers not found.")
        return text

    # Extract the course listings section
    course_section = text[start_of_section:end_of_section]

    # Dictionary to store courses and their years
    courses_years = {}

    # Find each year and its associated courses
    years_courses = re.findall(r'\\subsubsection\{(\d{4})\}(.*?)\\subsubsection|\\subsection', course_section, re.DOTALL)
    for year, courses in years_courses:
        for course_num in re.findall(r'(?:ME|M\s?E) (\d{3})', courses):
            if course_num not in excluded_courses:
                description = course_descriptions.get(course_num, f"ME {course_num}")
                if description not in courses_years:
                    courses_years[description] = []
                courses_years[description].append(year)

    # Format the new course section
    new_course_section = '\n\n' + r'\subsection{Teaching Experience}\label{teaching-experience}' + '\n'
    for course, years in courses_years.items():
        sorted_years = sorted(set(years), reverse=True)  # Sort and remove duplicates
        new_course_section += f"{course}, {', '.join(sorted_years)}.\n\n"

    # Replace the original course listings section with the new one
    new_text = text[:start_of_section] + new_course_section + text[end_of_section:]

    # Ensure a newline after \end{enumerate}
    new_text = new_text.replace('\end{enumerate}\subsection', '\end{enumerate}\n\n\subsection')

    return new_text

def replace_straight_quotes_with_latex_quotes(text_data):
    # This pattern matches phrases enclosed in straight double quotes, accounting for potential newlines
    pattern = r'"([\s\S]*?)"'
    
    # This function replaces straight double quotes with LaTeX quotes
    def replace_quotes(match):
        content = match.group(1)
        # Check if the content already starts with LaTeX-style quotes
        if content.startswith("``"):
            return f'"{content}"'  # Return the content as-is
        return f"``{content}''"
    
    # Use re.sub to replace all occurrences in the text
    corrected_text = re.sub(pattern, replace_quotes, text_data)
    
    # Debugging for the Service subsection
    service_section_pattern = r"\\subsection{Service}.*?\\subsubsection{Department}"
    service_section_match = re.search(service_section_pattern, corrected_text, re.DOTALL)
    
    return corrected_text

def clean_service_section(text_data):
    # Define subsubsection labels
    labels = ["College", "Department", "University", "Profession", "Society"]
    
    for label in labels:
        # Extract the subsubsection content
        section_pattern = rf"\\subsubsection{{{label}}}.*?(?=\\subsubsection|$)"
        section_match = re.search(section_pattern, text_data, re.DOTALL)
        
        if not section_match:
            print(f"{label} section not found in clean_service_section!")
            continue
        
        section_content = section_match.group(0)
        
        # Remove the redundant mention of the subsubsection title from each item
        cleaned_section = re.sub(rf"{label}, ", "", section_content)
        
        # Replace the old section with the cleaned section in the text data
        text_data = text_data.replace(section_content, cleaned_section)
    
    return text_data

def highlight_mentored_authors_astericks2(text_data):
    # Extract last names of mentored students and postdocs
    student_types = {
        'Ph.D. Dissertation': 'Ph.D. students',
        'Master\'s Thesis': 'Master students',
        'Postdoctoral Mentorship': 'Postdoc researchers',
        'Undergraduate Honors\nThesis': 'Undergrad students'
    }
    
    mentored_lastnames = []

    for subsection, label in student_types.items():
        start_index = text_data.find(rf'\subsubsection{{{subsection}}}')
        if start_index == -1:
            match = re.search(re.escape(subsection), text_data)
            if match:
                start_index = match.start()
            else:
                continue

        end_index = text_data.find(r'\subsubsection{', start_index + 1)
        if end_index == -1:
            end_index = len(text_data)

        section = text_data[start_index:end_index]
        pattern = r"\\item\s+(.*?),\s+``"
        matches = re.findall(pattern, section, re.DOTALL)
        
        for full_name in matches:
            lastname = full_name.split()[-1]
            mentored_lastnames.append(lastname)

    # Search through each publication and look for these last names
    start_publications = text_data.find(r'\subsection{Publications}')
    end_publications = text_data.find(r'\subsection{', start_publications + 1)
    if end_publications == -1:
        end_publications = len(text_data)

    publications_section = text_data[start_publications:end_publications]

    # Add a LaTeX asterisk after the name if it's found
    for lastname in mentored_lastnames:
        # Modified pattern to account for an optional middle initial
        pattern = rf"({lastname}(?:, [A-Z]\. ?[A-Z]?)?)"
        replacement = r"\1\\textsuperscript{*}"
        publications_section = re.sub(pattern, replacement, publications_section)

    # Replace the old publications section with the modified one
    modified_text = text_data[:start_publications] + publications_section + text_data[end_publications:]

    # Add mentored student note
    journal_article_section_start = modified_text.find(r'\subsubsection{Journal Article}\label{journal-article}')
    if journal_article_section_start == -1:
        print("Journal Article section not found in the text.")
        return modified_text

    note_text = "\n\\textit{Mentored student and postdoc co-authors have an astericks after their name.}\n"
    modified_text = modified_text[:journal_article_section_start] + note_text + modified_text[journal_article_section_start:]

    return modified_text

def underline_mentored_authors_with_note(text_data):
    # Extract last names of mentored students and postdocs
    student_types = {
        'Ph.D. Dissertation': 'Ph.D. students',
        'Master\'s Thesis': 'Master students',
        'Postdoctoral Mentorship': 'Postdoc researchers',
        'Undergraduate Honors\nThesis': 'Undergrad students'
    }
    
    mentored_lastnames = []

    for subsection, label in student_types.items():
        start_index = text_data.find(rf'\subsubsection{{{subsection}}}')
        if start_index == -1:
            match = re.search(re.escape(subsection), text_data)
            if match:
                start_index = match.start()
            else:
                continue

        end_index = text_data.find(r'\subsubsection{', start_index + 1)
        if end_index == -1:
            end_index = len(text_data)

        section = text_data[start_index:end_index]
        pattern = r"\\item\s+(.*?),\s+``"
        matches = re.findall(pattern, section, re.DOTALL)
        
        for full_name in matches:
            lastname = full_name.split()[-1]
            mentored_lastnames.append(lastname)

    # Search through each publication and look for these last names
    start_publications = text_data.find(r'\subsection{Publications}')
    end_publications = text_data.find(r'\subsection{', start_publications + 1)
    if end_publications == -1:
        end_publications = len(text_data)

    publications_section = text_data[start_publications:end_publications]

    # Add a LaTeX note about underlining mentored authors
    note_text = "\n\\textit{Mentored student and postdoc co-authors are underlined.}\n"
    journal_article_section_start = publications_section.find(r'\subsubsection{Journal Article}\label{journal-article}')
    publications_section = publications_section[:journal_article_section_start] + note_text + publications_section[journal_article_section_start:]

    # Underline the name if it's found
    for lastname in mentored_lastnames:
        # Modified pattern to account for an optional middle initial
        pattern = rf"({lastname}(?:, [A-Z]\. ?[A-Z]?)?)"
        replacement = r"\\underline{\1}"
        publications_section = re.sub(pattern, replacement, publications_section)

    # Replace the old publications section with the modified one
    modified_text = text_data[:start_publications] + publications_section + text_data[end_publications:]

    return modified_text

def create_custom_titles_for_sections(text_data):
    
    # change the title of the Jounral Article section
    text_data = text_data.replace(r'\subsubsection{Journal Article}\label{journal-article}', r'\subsubsection{Journal Articles}\label{journal-article}')
    
    # change the title of the Conference Proceeding section
    text_data = text_data.replace(r'\subsubsection{Conference Proceeding}\label{conference-proceeding}', r'\subsubsection{Conference Proceedings}\label{conference-proceeding}')
    
    # change the title of the Book Chapter section
    text_data = text_data.replace(r'\subsubsection{Book Chapter}\label{book-chapter}', r'\subsubsection{Book Chapters}\label{book-chapter}')

    # change the title of the Other section to Preprints and Technical Reports
    text_data = text_data.replace(r'\subsubsection{Other}\label{other}', r'\subsubsection{Preprints and Technical Reports}\label{other}')

    # in the Presentations section change the title of the Invited Talks section to Invited Talks and Seminars
    text_data = text_data.replace(r'\subsubsection{Invited}\label{invited}', r'\subsubsection{Invited Talks and Seminars}\label{invited}')

    # in the Presentations sections change the title of Uncategoried Presentations to Confererences and Workshops Presentations
    text_data = text_data.replace(r'\subsubsection{Uncategorized}\label{uncategorized}', r'\subsubsection{Conferences and Workshops}\label{uncategorized-presentations}')

    return text_data

def replace_section_colors(text, subsection_color='black', subsubsection_color='black'):
    # Define a pattern to search for the \titleformat command for \subsection and replace the color
    subsection_pattern = re.compile(r"(\\titleformat\{\\subsection\}.*?\\color\{)(blue)(\}.*)", re.DOTALL | re.VERBOSE)
    subsubsection_pattern = re.compile(r"(\\titleformat\{\\subsubsection\}.*?\\color\{)(red)(\}.*)", re.DOTALL | re.VERBOSE)

    # Define the replacement function for \subsection
    def subsection_replacement(match):
        #print("Subsection color found:", match.group(2))
        return f"{match.group(1)}{subsection_color}{match.group(3)}"

    # Define the replacement function for \subsubsection
    def subsubsection_replacement(match):
        #print("Subsubsection color found:", match.group(2))
        return f"{match.group(1)}{subsubsection_color}{match.group(3)}"

    # Perform the replacement for \subsection
    updated_text = subsection_pattern.sub(subsection_replacement, text)
    # Perform the replacement for \subsubsection
    updated_text = subsubsection_pattern.sub(subsubsection_replacement, updated_text)

    return updated_text

def process_professional_section(text):
    # Find the start of the "Professional Positions" section
    start_index = text.find("\\subsection{Professional Positions}")
    if start_index == -1:
        return "Error: Could not find the 'Professional Positions' section"
    
    # Find the end of the "Professional Positions" section
    end_index = text.find("\\subsection", start_index + 1)
    end_index = end_index if end_index != -1 else None
    
    # Extract the content before, within, and after the "Professional Positions" section
    before_professional_section = text[:start_index]
    professional_section = text[start_index:end_index]
    after_professional_section = text[end_index:] if end_index else ""
    
    # Normalize whitespace in the "Professional Positions" section
    professional_section = ' '.join(professional_section.split())

    # Example positions and locations
    positions_locations = [
        "Associate Professor of Biomedical Engineering (Courtesy), The Pennsylvania State University.",
        "Associate Professor of Mechanical Engineering, The Pennsylvania State University.",
        "Assistant Professor of Biomedical Engineering (Courtesy), The Pennsylvania State University.",
        "Assistant Professor of Mechanical Engineering, The Pennsylvania State University.",
        "Mechanical Engineer, The U.S. Army Research Laboratory, Soldier Protection Sciences Branch.",
        "Post-Doc, Oak Ridge Associated Universities at The U.S. Army Research Laboratory, Impact Physics Branch.",
        "Founder and Chief Engineer, BrainSim Technologies Inc.",
        "Lead Researcher of Computational Biomechanics, The Johns Hopkins University Applied Physics Laboratory, Research and Exploratory Development Department, Biomechanics and Injury Mitigation Systems Group."
    ]

    # Process each position and location
    for pl in positions_locations:
        normalized_pl = ' '.join(pl.split())
        
        # Find the position and location in the content
        position_start = professional_section.find(normalized_pl)
        if position_start == -1:
            print("Error: Could not find position and location:", pl)
            continue
        
        # Extract the position, location, and date info
        position_end = position_start + len(normalized_pl)
        rest_of_content = professional_section[position_end:]
        date_match = re.search(r'\((.*?)\)', rest_of_content)
        
        if not date_match:
            print("Error: Could not find date information for:", pl)
            continue
        
        date_info = date_match.group(1)
        updated_date_info = process_dates(date_info)
        
        # Replace old date string with new date format
        date_start = position_end + rest_of_content.find(date_info)
        date_end = date_start + len(date_info) + 2  # +2 to include parentheses
        updated_entry = normalized_pl + " (" + updated_date_info + "). "
        updated_entry += "\n"


        professional_section = professional_section[:position_start] + updated_entry + professional_section[date_end:]
        
    # Reconstruct the entire text string with the updated "Professional Positions" section
    updated_text = before_professional_section + professional_section + after_professional_section
    
    # Add line returns for better readability in the .tex source file
    updated_text = updated_text.replace('. ', '. \n')

    # Add line returns after specific subsections and subsubsections
    updated_text = updated_text.replace("\\subsection{Professional Positions}\label{professional-positions}", "\\subsection{Professional Positions}\label{professional-positions}\n")
    updated_text = updated_text.replace("\\subsubsection{Academic}\label{academic}", "\\subsubsection{Academic}\label{academic}\n")


    return updated_text


def process_dates(date_string):
    # Define possible date formats
    date_formats = ["%B %Y", "%B %d, %Y"]

    # Define the function for processing dates
    def process_single_date(date):
        for fmt in date_formats:
            try:
                return datetime.datetime.strptime(date, fmt).strftime("%Y")
            except ValueError:
                continue
        return "Present" if "Present" in date else "Invalid Date"

    date_ranges = date_string.split(" - ")
    start_date = process_single_date(date_ranges[0])
    end_date = process_single_date(date_ranges[1]) if len(date_ranges) > 1 else "Present"

    return f"{start_date} - {end_date}"
def format_positions(text):
    # Define the pattern to match the positions and dates
    pattern = re.compile(r'(.*?), The (.*?)\.\s+\((.*?)\)\.')

    def replace_with_tabular(match):
        # Extract the position, location, and date
        position = match.group(1).strip()
        location = "The " + match.group(2).strip()
        date = match.group(3).strip()

        # Return the formatted tabular environment string
        return (
            r"\begin{tabular}{@{}l@{\hskip 0.5in}r@{}}"
            f"{position}, {location}. & ({date}). \\\\"
            r"\end{tabular}"
        )

    # Search and replace all occurrences of the pattern with the tabular format
    updated_text = re.sub(pattern, replace_with_tabular, text)

    return updated_text


def format_professional_positions_section(text, spacing='5pt'):
    # Add the booktabs package if not present
    if '\\usepackage{booktabs}' not in text:
        text = text.replace(r'\usepackage{xcolor}', r'\usepackage{xcolor}' + '\n' + r'\usepackage{booktabs}')

    # Insert the tabularx package right after the xcolor package if not present
    if '\\usepackage{tabularx}' not in text:
        text = text.replace(r'\usepackage{xcolor}', r'\usepackage{xcolor}' + '\n' + r'\usepackage{tabularx}')

    # Pattern to extract the Professional Positions section
    positions_section_pattern = r'(\\subsection\{Professional Positions\}\\label\{professional-positions\})(.+?)(\\subsection\{Education\}\\label\{education\})'
    match = re.search(positions_section_pattern, text, re.DOTALL)

    if match:
        # Extract the subsection title and the 'Education' section header
        subsection_title = match.group(1).strip()
        education_section_header = match.group(3).strip()

        # Initialize an empty string to store the formatted Professional Positions section
        formatted_positions = subsection_title + "\n\n"

        professional_positions_section = match.group(2)

        # Split the section into subsubsections, capturing the titles
        subsubsections = re.split(r'(\\subsubsection\{.*?\}\\label\{.*?\})', professional_positions_section)

        # Process each subsubsection and its entries
        for i in range(1, len(subsubsections), 2):
            subsubsection_title = subsubsections[i].strip()
            formatted_positions += subsubsection_title + "\n\n"

            # Begin the tabularx environment for the subsubsection
            formatted_positions += "\\begin{tabularx}{\\textwidth}{@{}Xr@{}}\n"

            # Split into entries
            entries = subsubsections[i + 1].strip().split('\n\n')
            for index, entry in enumerate(entries):
                entry = entry.strip()
                if entry:
                    # Use regular expression to separate the position/location from the date range
                    date_range_pattern = r'\(\d{4} - (?:\d{4}|Present)\)'
                    date_range_match = re.search(date_range_pattern, entry)
                    if date_range_match:
                        date_range = date_range_match.group(0).strip('()')  # Remove parentheses
                        pos_loc = entry[:date_range_match.start()].strip()
                        formatted_positions += f"{pos_loc} & {date_range}"
                        # Add the line break and space unless it's the last entry
                        if index < len(entries) - 1:
                            formatted_positions += f" \\\\\n\\addlinespace[{spacing}]\n"
                    else:
                        # Handle the case where the date range is not found
                        formatted_positions += f"{entry} & N/A"
                        # Add the line break and space unless it's the last entry
                        if index < len(entries) - 1:
                            formatted_positions += f" \\\\\n\\addlinespace[{spacing}]\n"

            # End the tabularx environment for the subsubsection
            formatted_positions += "\\end{tabularx}\n\n"

        # Replace the original Professional Positions section with the formatted section
        text = text.replace(match.group(0), formatted_positions + education_section_header)

    else:
        print("The 'Professional Positions' subsection was not found.")

    return text

def format_education_section(text, spacing='5pt'):
    # Define the pattern to find the education section
    education_section_pattern = r'(\\subsection\{Education\}\\label\{education\})(.+?)(\\subsection\{.*?\}\\label\{.*?\}|\Z)'
    match = re.search(education_section_pattern, text, re.DOTALL)

    if match:
        subsection_title = match.group(1).strip()
        following_section_header = match.group(3).strip()
        education_entries = match.group(2).strip().split('\n\n')

        # Initialize an empty string to store the formatted Education section
        formatted_education = subsection_title + "\n\n\\begin{flushleft}\n\\begin{tabularx}{\\textwidth}{@{}Xr@{}}\n"

        # Process each entry
        for index, entry in enumerate(education_entries):
            if entry:
                # Remove any trailing period from each entry before the last line break
                entry = re.sub(r'\.\s*(?=\\|$)', '', entry)
                # Split entry into institution/major and year
                parts = entry.rsplit(',', 1)
                if len(parts) == 2:
                    institution_major = parts[0].strip()
                    year = parts[1].strip().lstrip('\\').rstrip('.')
                    formatted_education += f"{institution_major} & {year}"
                else:
                    formatted_education += f"{entry} & N/A"

                # Add the line break and space unless it's the last entry
                if index < len(education_entries) - 1:
                    formatted_education += f" \\\\\n\\addlinespace[{spacing}]\n"

        # Close the tabularx and flushleft environments
        formatted_education += "\n\\end{tabularx}\n\\end{flushleft}\n\n" + following_section_header

        # Replace the original Education section with the formatted section
        text = text.replace(match.group(0), formatted_education)

    else:
        print("The 'Education' subsection was not found.")

    return text

def format_awards_and_honors_section(text, spacing='5pt', borders=False):
    # Add the longtable package if not present
    if '\\usepackage{longtable}' not in text:
        text = text.replace(r'\usepackage{tabularx}', r'\usepackage{tabularx}' + '\n' + r'\usepackage{longtable}')

    # Pattern to extract the Awards and Honors section
    awards_section_pattern = r'(\\subsection\{Awards and Honors\}\\label\{awards-and-honors\})(.+?)(\\subsection\{Publications\}\\label\{publications\})'
    match = re.search(awards_section_pattern, text, re.DOTALL)

    if match:
        # Extract the subsection title and the 'Publications' section header
        subsection_title = match.group(1).strip()
        publications_section_header = match.group(3).strip()

        # Initialize an empty string to store the formatted Awards and Honors section
        formatted_awards = subsection_title + "\n\n"

        # Define new lengths for table columns
        #formatted_awards += "\\newlength{\\mylength}\n"
        #formatted_awards += "\\newlength{\\myotherlength}\n"
        #formatted_awards += "\\setlength{\\mylength}{\\dimexpr(\\textwidth-0\\tabcolsep)*3/4\\relax}\n"
        #formatted_awards += "\\setlength{\\myotherlength}{\\dimexpr(\\textwidth-2\\tabcolsep)/4\\relax}\n"
        
        # Define the fixed width for the second column
        fixed_second_column_width = "1.0in"  # or the width you prefer

        # Modify the formatted_awards string with the new lengths
        #formatted_awards = ""
        formatted_awards += "\\newlength{\\mylength}\n"
        formatted_awards += "\\newlength{\\myotherlength}\n"
        formatted_awards += f"\\setlength{{\\myotherlength}}{{{fixed_second_column_width}}}\n"  # Set the fixed width for the second column
        formatted_awards += "\\setlength{\\mylength}{\\dimexpr(\\textwidth-\\myotherlength-0\\tabcolsep)\\relax}\n"  # Calculate the remaining width for the first column

        # Define the table format string based on whether borders are enabled
        if borders:
            table_format = "|p{\\mylength}|>{\\raggedleft\\arraybackslash}p{\\myotherlength}|"
        else:
            table_format = "@{}p{\\mylength}@{}>{\\raggedleft\\arraybackslash}p{\\myotherlength}@{}"

        # Begin the longtable environment for the section
        formatted_awards += f"\\begin{{longtable}}{{{table_format}}}\n"
        if borders:
            formatted_awards += "\\hline\n"

        awards_and_honors_section = match.group(2).strip()

        # Split into entries by looking for two consecutive newlines
        entries = awards_and_honors_section.strip().split('\n\n')
        for index, entry in enumerate(entries):
            entry = entry.strip()
            if entry:
                # Use regular expression to separate the award/honor description from the date
                date_pattern = r'\((?:\w+ )?(\d{4})(?: - (\d{4}|\w+))?\.?\)'
                entry = re.sub(r'\s+', ' ', entry)  # Replace multiple whitespaces with a single space
                date_match = re.search(date_pattern, entry)
                if date_match:
                    # Extract the year or year range
                    year = date_match.group(1)
                    award_desc = entry[:date_match.start()].strip()
                    # Add vertical space before each entry except the first one
                    if index > 0:
                        formatted_awards += "\n\\addlinespace[{}]\n".format(spacing)
                    formatted_awards += f"{award_desc} & {year} \\\\"
                    if borders and index < len(entries) - 1:
                        formatted_awards += "\\hline\n"
                else:
                    # Handle the case where the date is not found
                    if index > 0:
                        formatted_awards += "\n\\addlinespace[{}]\n".format(spacing)
                    formatted_awards += f"{entry} & N/A \\\\"
                    if borders and index < len(entries) - 1:
                        formatted_awards += "\\hline\n"

        # End the longtable environment for the section
        formatted_awards += "\\end{longtable}\n"

        # Replace the original Awards and Honors section with the formatted section
        text = text.replace(match.group(0), formatted_awards + publications_section_header)

    else:
        print("The 'Awards and Honors' subsection was not found.")

    return text

def update_documentclass_font_size(latex_content, font_size):
    # Define the pattern to find the documentclass declaration
    pattern = re.compile(r'\\documentclass\[(.*?)\]\{article\}', re.DOTALL)
    
    # Define the replacement string to include the desired font size option
    replacement = fr'\\documentclass[{font_size}pt]{{article}}'
    
    # Replace the documentclass declaration with the one including the desired font size option
    updated_content = re.sub(pattern, replacement, latex_content)
    
    return updated_content

def main():
    #open a file to read in C:\Users\rhk12\OneDrive - The Pennsylvania State University\resume\CV named main.tex 
    # open the file for reading
    #f = open(r'C:\Users\rhk12\OneDrive - The Pennsylvania State University\resume\CV\main.tex', 'r')
    #f = open(r'C:\Users\rhk12\OneDrive - The Pennsylvania State University\resume\CV\main.tex', 'r', encoding='utf-8')
    f = open(r'C:\Users\rhk12\OneDrive - The Pennsylvania State University\resume\CV\trypandoc11162023.tex', 'r', encoding='utf-8')

    #open another file for writing
    f1 = open(r'C:\Users\rhk12\OneDrive - The Pennsylvania State University\resume\CV\main_edited.tex', 'w')

    #dossier_file_path = r"C:\Users\rhk12\OneDrive - The Pennsylvania State University\resume\CV\20231021-095357-CDT.docx"
    dossier_file_path = r"C:\Users\rhk12\OneDrive - The Pennsylvania State University\resume\CV\20231024-203540-CDT.docx"
    
    #read in the entire file and store it in a variable called text 
    text = f.read()
        
    # add my formating package
    text = add_custom_package(text)
    
    # Set the section colors
    text = set_section_colors(text)
      
     # Format the header
    text = format_header(text)
    
    # add the date to the latex file
    text = add_date_to_header(text)

    # write_courses_to_file(text, f1)
    text = process_courses(text)
    
    # reorder the publications
    text = reorder_publications(text)
    
    # boldface my name in the publications
    text = boldface_name_in_publications(text)
        
    # reorder the student sections
    text = reorder_student_sections4(text)
    
    # process the student thesis titles
    text = process_student_thesis_titles(text,dossier_file_path)

    # clean up the service section to remove extra text
    text = clean_service_section(text)

    # replace straight quotes with latex quotes
    text = replace_straight_quotes_with_latex_quotes(text)
    
    # add emphasis for mentees  - pick between astericks or underlines
    #text = highlight_mentored_authors_astericks2(text)
    text = underline_mentored_authors_with_note(text)
    
    # format the professional section   
    text = process_professional_section(text)
    
    # convert the professional positions section to a tabularx environment
    # so the date range can be right justified
    text = format_professional_positions_section(text,spacing='5pt')

    # convert the education section to a tabularx environment
    # so the date range can be right justified
    text = format_education_section(text, spacing='5pt')
    
    # format the awards and honors section
    # so the date range can be right justified
    text = format_awards_and_honors_section(text, spacing='5pt', borders=False)
    
    # ------------------------------------------
    # --------- Final Formatting ----------------
    # ------------------------------------------

    # add custom titles for sections
    text = create_custom_titles_for_sections(text)
    
    # colorize the subsection and subsubsection titles
    text = replace_section_colors(text, subsection_color='black', subsubsection_color='black')
    
    # make subsection text uppercase - needs to be last
    text = capitalize_subsections(text)
    
    text = update_documentclass_font_size(text, 11)

  
    f1.write(text)
    
    #close the files
    f.close()
    f1.close()
    

if __name__ == "__main__":
    main()

    