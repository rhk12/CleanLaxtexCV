#create a python script to clean the latex cv

# import the necessary packages
import argparse
import math
import sys
import re
import shutil
import subprocess
import time
import glob
import csv
import random
import string
import datetime
import os 
import datetime
from docx import Document 
from collections import defaultdict
import unidecode


def process_student_thesis_titles(text): 
    
    # Extract PhD dissertation titles
    phd_pattern = r'\\item\s+([^,]+,\s+[A-Z]\.),\s+Ph\.D\.\s+(.+?)(?:\.\s+\(.*?\))?\.'
    phd_data = extract_titles(text, "Ph\\.D\\. Dissertation Advisor", phd_pattern)
    print(phd_data)
    # Extract postdoc titles
    postdoc_pattern = r'\\item\s+([^,]+,\s+[A-Z]\.?),\s+([A-Z])\.\s+(.+?)(?:\.\s+\(.*?\))?\.'
    postdoc_data = extract_titles(text, "Postdoctoral Mentorship Advisor", postdoc_pattern)

    # Extract undergrad thesis titles
    undergrad_pattern = r'\\item\s+([^,]+,\s+[A-Z]\.?),\s+Undergraduate\.\s+(.+?)(?:\.\s+\(.*?\))?\.'
    undergrad_data = extract_titles(text, "Undergraduate Honors Thesis Advisor", undergrad_pattern)
     
    # Merge the dictionaries
    student_data = { **phd_data,**postdoc_data, **undergrad_data}

    # Replace problematic characters in titles
    student_data = replace_problematic_characters_in_titles(student_data)
      
    # Update the LaTeX content with thesis titles - does not work for postdoc titles see next function
    text = add_title_to_name2(text, student_data)
       
    # reformat phd section for name, title, date
    text = reformat_phd_section(text)
    
    # reformat masters section for name, title, date    
    text = update_masters_section(text)
    
    # reformat postdoc section for name, title, date
    text = update_postdoc_section(text)
    
    # reformat undergrad section for name, title, date
    text = update_undergraduate_section(text)
   
    return text


def replace_problematic_characters_in_titles(student_data):
    problematic_char = u'\u2013'  # En dash (U+2013)

    for name, title in student_data.items():
        if problematic_char in title:
            #print(f"Found character '{problematic_char}' in {name}: {title}")
            student_data[name] = title.replace(problematic_char, '--')

    return student_data

def update_undergraduate_section(text_data):
    # Define section label
    label = "Undergraduate Honors\nThesis"
    
    # Extract the Undergraduate Honors Thesis section
    section_pattern = rf"\\subsubsection{{{label}}}.*?\\begin{{enumerate}}.*?\\end{{enumerate}}"
    section_match = re.search(section_pattern, text_data, re.DOTALL)
    
    if not section_match:
        print("Undergraduate Honors Thesis section not found!")
        return text_data
    
    section_content = section_match.group(0)
    
    # Extract individual items
    item_pattern = r"Undergraduate Honors Thesis\. \((.*?)\).\\\\\s+Advised: (.*?), \"(.*?)\","
    items = re.findall(item_pattern, section_content, re.DOTALL)
    
    # Format items
    formatted_items = []
    for time_period, name, title in items:
        formatted_item = f"\\item {name}, \"{title}\", {time_period}"
        formatted_items.append(formatted_item)
    
    # Combine formatted items
    formatted_section = "\\subsubsection{Undergraduate Honors\nThesis}\n\\begin{enumerate}\n\\def\\labelenumi{\\arabic{enumi}.}\n" + "\n".join(formatted_items) + "\n\\end{enumerate}"
    
    # Replace the old section with the new formatted section in the text data
    updated_text_data = text_data.replace(section_content, formatted_section)
    
    return updated_text_data

def update_postdoc_section(text_data):
    # Define section label
    label = "Postdoctoral Mentorship"
    
    # Extract the Postdoctoral Mentorship section
    section_pattern = rf"\\subsubsection{{{label}}}.*?\\begin{{enumerate}}.*?\\end{{enumerate}}"
    section_match = re.search(section_pattern, text_data, re.DOTALL)
    
    if not section_match:
        print("Postdoctoral Mentorship section not found!")
        return text_data
    
    section_content = section_match.group(0)
    
    # Extract individual items
    item_pattern = r"Postdoctoral Mentorship\. \((.*?)\).\\\\\s+Advised: (.*?), \"(.*?)\","
    items = re.findall(item_pattern, section_content, re.DOTALL)
    
    # Format items
    formatted_items = []
    for time_period, name, title in items:
        formatted_item = f"\\item {name}, \"{title}\", {time_period}"
        formatted_items.append(formatted_item)
    
    # Combine formatted items
    formatted_section = "\\subsubsection{Postdoctoral Mentorship}\n\\begin{enumerate}\n\\def\\labelenumi{\\arabic{enumi}.}\n" + "\n".join(formatted_items) + "\n\\end{enumerate}"
    
    # Replace the old section with the new formatted section in the text data
    updated_text_data = text_data.replace(section_content, formatted_section)
    
    return updated_text_data

def update_masters_section(text_data):
    # Define the two possible formats for "Master's Thesis"
    master_thesis_options = ["Master's Thesis", "Master\\textquotesingle s Thesis"]

    # Detect which format is used in the text
    master_thesis_format = next((option for option in master_thesis_options if option in text_data), None)

    if master_thesis_format is None:
        print("Master's Thesis section format not detected!")
        return text_data

    # Escape backslashes for regex
    escaped_master_thesis_format = master_thesis_format.replace('\\', '\\\\')

    # Prepare regex pattern for the section
    section_pattern = rf"(\\subsubsection\{{.*?{escaped_master_thesis_format}.*?\}}.*?\\begin\{{enumerate\}})(.*?)(\\end\{{enumerate\}})"
    
    section_match = re.search(section_pattern, text_data, re.DOTALL)

    if not section_match:
        print("Master's Thesis section not found!")
        return text_data

    section_start, section_content, section_end = section_match.groups()
    
    # Adjusting item pattern for matching items within the section
    item_pattern = rf"\\item\s*{escaped_master_thesis_format}\.\s*\((.*?)\)\.\\\\\s*Advised:\s*(.*?),\s*\"(.*?)\""
    
    items = re.findall(item_pattern, section_content, re.DOTALL | re.MULTILINE)
    
    if not items:
        print("No items found in Master's Thesis section!")
        return text_data

    # Format items
    formatted_items = [f"\\item {name}, \"{title}\", {time_period}" for time_period, name, title in items]

    # Combine formatted items
    formatted_section_content = "\n".join(formatted_items)

    # Replace the old section with the new formatted section in the text data
    updated_section = section_start + formatted_section_content + section_end
    updated_text_data = text_data.replace(section_content, formatted_section_content)
    
    return updated_text_data

def reformat_masters_section(text_data):
    # Define section label
    label = "Master's Thesis"
    
    # Extract the Master's Thesis section
    section_pattern = rf"(\\subsubsection{{{label}.*?}}.*?\\begin{enumerate}.*?\\end{enumerate})"
    section_match = re.search(section_pattern, text_data, re.DOTALL)
    
    if not section_match:
        return text_data
    
    section_content = section_match.group(1)
    
    # Pattern to extract student details
    student_pattern = rf"{label}. \((.*?)\).+?Advised: (.*?), \"(.*?)\","
    student_matches = re.findall(student_pattern, section_content, re.DOTALL)
    
    # Create a new formatted section
    new_section = f"\\subsubsection{{{label}}}\n\n\\begin{enumerate}\n"
    for time_period, name, title in student_matches:
        # Split name into first and last names
        names = name.split()
        first_name = names[0]
        last_name = names[-1]
        
        # Add to the new section
        new_section += f"\\item\n  {first_name} {last_name}, \"{title}\", {time_period}\n"
    
    new_section += "\\end{enumerate}\n"
    
    # Replace the old section with the new one
    text_data = text_data.replace(section_content, new_section)
    
    return text_data

def reformat_sections(text_data):
    # Define section labels
    section_labels = {
        "postdoc": "Postdoctoral Mentorship",
        "undergrad": "Undergraduate Honors Thesis",
        "masters": "Master's Thesis"
    }
    
    # Function to reformat a specific section
    def reformat_section(section_content, label):
        student_pattern = rf"{label}. \((.*?)\).+?Advised: (.*?), \"(.*?)\","
        student_matches = re.findall(student_pattern, section_content, re.DOTALL)
        
        # Create a new formatted section
        new_section = f"\\subsubsection{{{label}}}\n\n\\begin{enumerate}\n"
        for time_period, name, title in student_matches:
            # Split name into first and last names
            names = name.split()
            first_name = names[0]
            last_name = names[-1]
            
            # Add to the new section
            new_section += f"\\item\n  {first_name} {last_name}, \"{title}\", {time_period}\n"
        
        new_section += "\\end{enumerate}\n"
        return new_section
    
    # Iterate over each section label and reformat
    for section, label in section_labels.items():
        section_pattern = rf"(\\subsubsection{{{label}.*?}}.*?\\begin{enumerate}.*?\\end{enumerate})"
        section_match = re.search(section_pattern, text_data, re.DOTALL)
        
        if section_match:
            section_content = section_match.group(1)
            new_section = reformat_section(section_content, label)
            text_data = text_data.replace(section_content, new_section)
    
    return text_data

def extract_undergrad_student_titles(file_path):
    #file_path = r"C:\Users\rhk12\OneDrive - The Pennsylvania State University\resume\CV\20231021-095357-CDT.docx"

    # Check if the file exists
    if not os.path.exists(file_path):
        return f"File not found at '{file_path}'"
    
    # If the file exists, proceed to read its content
    doc = Document(file_path)
    full_text = [para.text for para in doc.paragraphs]

    # Define the sub-section titles
    sub_section_title = "Undergraduate Honors Thesis Advisor"
    end_section_title = "THE SCHOLARSHIP OF Research and Creative Accomplishments"

    # Find the start of the sub-section
    start_index = next((i for i, text in enumerate(full_text) if sub_section_title in text), None)
    if start_index is None:
        return {}

    # Find the end of the sub-section
    end_index = next((i for i, text in enumerate(full_text[start_index:]) if end_section_title in text), None)
    if end_index is None:
        end_index = len(full_text)
    else:
        end_index += start_index

    # Extract undergrad student names and thesis titles
    undergrad_data = {}
    for line in full_text[start_index + 1:end_index]:
        if ", Undergraduate." in line:
            name, title_with_dates = line.split(", Undergraduate.", 1)
            title = title_with_dates.split(". (")[0].strip()
            
            # Remove the "Date Graduated" part from the title
            title = re.sub(r"\. Date Graduated:.*$", "", title)
            
            undergrad_data[name.strip()] = title
    

    return undergrad_data

def add_postdoc_work(text, postdoc_data):
    for name, title in postdoc_data.items():
        # Split the name into first and last names
        last_name, first_initial = name.split(', ')
        first_initial = first_initial.strip()[0]  # Get the first character of the first name
        
        # Construct two patterns: one for full name and one for first initial
        full_name_pattern = rf"(\\subsubsection\{{Postdoctoral Mentorship\}}.*?Advised: ){first_initial}.*?{last_name}"
        initial_pattern = rf"(\\subsubsection\{{Postdoctoral Mentorship\}}.*?Advised: ){first_initial} {last_name}"
        
        # Try matching with the full name
        if re.search(full_name_pattern, text, flags=re.DOTALL):
            text = re.sub(full_name_pattern, f"\\1{first_initial}. {last_name}, \"{title}\"", text, flags=re.DOTALL)
            print(f"Updated {first_initial}. {last_name}'s entry.")  # Debugging print statement
        # If that doesn't work, try matching with the first initial
        elif re.search(initial_pattern, text, flags=re.DOTALL):
            text = re.sub(initial_pattern, f"\\1{first_initial}. {last_name}, \"{title}\"", text, flags=re.DOTALL)
            print(f"Updated {first_initial}. {last_name}'s entry using initial.")  # Debugging print statement
        else:
            print(f"Pattern not found for: {first_initial}. {last_name} or {first_initial} {last_name}")  # Debugging print statement

    return text

def reformat_phd_section(text_data):
    # Extract the Ph.D. section
    phd_section_pattern = r"(\\subsubsection{Ph\.D\. Dissertation}.*?\\begin{enumerate}.*?\\end{enumerate})"
    phd_section_match = re.search(phd_section_pattern, text_data, re.DOTALL)
    
    if not phd_section_match:
        print("Ph.D. section not found in the provided text.")
        return text_data
    
    phd_section = phd_section_match.group(1)
    
    # Extract details for each student
    student_pattern = r"Ph\.D\. Dissertation\. \((.*?)\).+?Advised: (.*?), \"(.*?)\","
    student_matches = re.findall(student_pattern, phd_section, re.DOTALL)
    
    # Create a new formatted Ph.D. section
    new_phd_section = "\\subsubsection{Ph.D. Dissertation}\n\n\\begin{enumerate}\n"
    for time_period, name, title in student_matches:
        # Split name into first and last names
        names = name.split()
        first_name = names[0]
        rest_of_name = " ".join(names[1:])
        
        # Add to the new Ph.D. section
        new_phd_section += f"\\item\n  {first_name} {rest_of_name}, \"{title}\", {time_period}\n"

    
    new_phd_section += "\\end{enumerate}\n"
    
    # Replace the original Ph.D. section with the new one
    updated_text_data = text_data.replace(phd_section, new_phd_section)
    
    
    return updated_text_data


def add_title_to_name2(text, data):
    # Extract the section between \subsection{Directed Student Learning} and \subsection{Teaching Experience}
    start_index = text.find(r'\subsection{Directed Student Learning}')
    end_index = text.find(r'\subsection{Teaching Experience}', start_index)
    
    if start_index == -1 or end_index == -1:
        print("Directed Student Learning or Teaching Experience section not found!")
        return text

    section = text[start_index:end_index]

    for name, title in data.items():
        # Adjust for the name format in the postdoc section
        if ',' in name:
            last_name, first_initial = name.split(', ')
            formatted_name_full = f"{first_initial}. {last_name}"
            formatted_name_last = last_name
        else:
            formatted_name_full = name
            formatted_name_last = name.split()[0]  # Just the first name
        
        # Check which naming convention exists in the section
        if formatted_name_full in section:
            pattern = re.escape(formatted_name_full)
        elif formatted_name_last in section:
            pattern = re.escape(formatted_name_last)
        else:
            print(f"Pattern not found for: {name}")
            continue
        
        # Extract the time period
        time_period_match = re.search(r"\((.*?)\)", section)
        if time_period_match:
            time_period = time_period_match.group(1)
        else:
            time_period = ""

        # If we're dealing with the postdoc section
        if "Postdoctoral Mentorship" in title:
            postdoc_section_pattern = r"(\\subsubsection\{Postdoctoral Mentorship\}.*?)" + pattern
            match = re.search(postdoc_section_pattern, section, re.DOTALL)
            if match:
                section = section.replace(match.group(0), match.group(0).replace(f"Advised: {formatted_name_full}", f"{formatted_name_last}, \"{title}\", {time_period}"))
        # If we're dealing with the undergraduate section
        elif "Undergraduate Honors Thesis" in title:
            undergrad_section_pattern = r"(\\subsubsection\{Undergraduate Honors\s*Thesis\}.*?Advised: )" + pattern
            match = re.search(undergrad_section_pattern, section, re.DOTALL)
            if match:
                section = section.replace(match.group(0), match.group(0).replace(f"Advised: {formatted_name_full}", f"{formatted_name_last}, \"{title}\", {time_period}"))
        # For Master's and Ph.D. students
        else:
            if re.search(pattern, section):
                section = re.sub(pattern, f"{formatted_name_last}, \"{title}\", {time_period}", section)
    # Define patterns for subsections to remove
    sections_to_remove = [
        r"\\subsection\{Postdoctoral Mentorship Advisor\}.*?(?=\\subsection|\Z)",
        r"\\subsection\{Ph\.D\. Dissertation Committee Member\}.*?(?=\\subsection|\Z)",
        r"\\subsection\{Research Activity Advisor\}.*?(?=\\subsection|\Z)",
        r"\\subsection\{Undergraduate Honors Thesis Advisor\}.*?(?=\\subsection|\Z)",
        r"\\subsection\{Ph.D. Dissertation Advisor\}.*?(?=\\subsection|\Z)"
    ]

    # Remove specified subsections
    for pattern in sections_to_remove:
        text = re.sub(pattern, '', text, flags=re.DOTALL)
    
    # Replace the original section with the modified one
    text = text[:start_index] + section + text[end_index:]
    # Ensure \end{document} is present at the end of the document
    if not text.strip().endswith("\\end{document}"):
        text += "\n\\end{document}"
    return text

def add_undergrad_titles(text, undergrad_data):
    for name, title in undergrad_data.items():
        # Convert the name format from 'Last, F.' to 'First Last'
        last_name, first_initial = name.split(', ')
        # We'll use a placeholder for the first name since we don't have the full first name
        first_name_placeholder = f"{first_initial}.*"
        
        # Define the pattern to search for the student's entry in the LaTeX content
        pattern = rf"(Advised: {first_name_placeholder} {last_name})"
        
        # Check if the pattern exists in the text
        if re.search(pattern, text):
            # Replace the pattern with the pattern + title
            text =  re.sub(pattern, f"\\1, \"{title}\"", text)
            print(f"Updated {first_initial}. {last_name}'s entry.")
        else:
            print(f"Pattern for {first_initial}. {last_name} not found in the text!")
    
    return text

def extract_titles_from_word(doc, section_name, entry_pattern):
    # Initialize a dictionary to store the extracted data
    extracted_data = {}
    
    # Locate the section
    section_found = False
    section_content = []
    
    for para in doc.paragraphs:
        if section_name in para.text:
            section_found = True
        elif section_found and para.text.startswith('Section'):
            break
        elif section_found:
            section_content.append(para.text)
    
    if section_content:
        content_str = "\n".join(section_content).strip()
        entries = re.findall(entry_pattern, content_str, re.MULTILINE)

        for entry in entries:
            name, title = entry
            extracted_data[name.strip()] = title.strip()
    else:
        extracted_data["error"] = f"{section_name} section not found in the provided document."

    return extracted_data


def extract_titles(tex_content, section_name, entry_pattern):

    # Pattern to extract the specified section
    pattern = rf'\\subsection\{{{section_name}\}}(.*?)\\end\{{enumerate\}}'
    match = re.search(pattern, tex_content, re.DOTALL)

    # Initialize a dictionary to store the extracted data
    extracted_data = {}

    if match:
        
        if section_name=="Master's Thesis":

            section_content = match.group(0).strip()
            
        else:
            section_content = match.group(1).strip()
        # Extract individual entries
        entries = re.findall(entry_pattern, section_content, re.MULTILINE)

        # Populate the dictionary with extracted data
        for entry in entries:
            name, title = entry
            extracted_data[name.strip()] = title.strip()
    else:
        extracted_data["error"] = f"{section_name} section not found in the provided text."

    return extracted_data

def extract_subsubsection(text, title):
    # Use regex to extract the subsubsection based on the title
    pattern = re.compile(r'\\subsubsection{' + re.escape(title) + r'}.*?\\end{enumerate}', re.DOTALL)
    match = pattern.search(text)
    if match:
        return match.group(0)
    else:
        print(f"Title '{title}' not found in extract_subsubsection!")
        return None

def reorder_student_sections(text):
    # Define the desired order
    order = ["Ph.D. Dissertation", "Master's Thesis", "Postdoctoral Mentorship", "Undergraduate Honors\nThesis"]
    
    # Extract each subsubsection based on the order
    sections = [extract_subsubsection(text, title) for title in order]
    
    # Check if any section is missing and remove it
    sections = [section for section in sections if section is not None]
    
    # Concatenate the reordered sections
    reordered_sections = "\n".join(sections)
    
    # Replace the original Directed Student Learning section with the reordered one
    start_idx = text.find(r'\subsection{Directed Student Learning}')
    end_idx = text.find(r'\subsection', start_idx + 1)
    if end_idx == -1:
        end_idx = len(text)
    
    return text[:start_idx] + r'\subsection{Directed Student Learning}' + '\n' + reordered_sections + text[end_idx:]
def reorder_student_sections4(text):
    # Define the two possible formats for "Master's Thesis"
    master_thesis_options = ["Master's Thesis", "Master\\textquotesingle s Thesis"]

    # Detect which format is used in the text
    master_thesis_format = next((option for option in master_thesis_options if option in text), None)

    if master_thesis_format is None:
        print("Master's Thesis section format not detected")
        return text

    # Define the desired order based on the detected format
    order = ["Ph.D. Dissertation", master_thesis_format, "Postdoctoral Mentorship", "Undergraduate Honors\nThesis"]

    # Extract each subsubsection based on the order
    sections = [extract_subsubsection(text, title) for title in order]

    # Check if any section is missing and remove it
    sections = [section for section in sections if section is not None]

    # Concatenate the reordered sections
    reordered_sections = "\n".join(sections)

    # Replace the original Directed Student Learning section with the reordered one
    start_idx = text.find(r'\subsection{Directed Student Learning}')
    end_idx = text.find(r'\subsection', start_idx + 1)
    if end_idx == -1:
        end_idx = len(text)
    #print("aaa",reordered_sections)
    return text[:start_idx] + r'\subsection{Directed Student Learning}' + '\n' + reordered_sections + text[end_idx:]
def reorder_student_sections3(text):
    # Define the two possible orders for "Master's Thesis"
    order1 = ["Ph.D. Dissertation", "Master's Thesis", "Postdoctoral Mentorship", "Undergraduate Honors\nThesis"]
    order2 = ["Ph.D. Dissertation", "Master\\textquotesingle s Thesis", "Postdoctoral Mentorship", "Undergraduate Honors\nThesis"]

    # Extract each subsubsection based on the two orders
    sections1 = [extract_subsubsection(text, title) for title in order1]
    sections2 = [extract_subsubsection(text, title) for title in order2]

    # Combine the two sets of sections, removing any None entries
    combined_sections = [section for section in sections1 + sections2 if section is not None]

    # Remove duplicate sections (if any)
    seen_titles = set()
    unique_sections = []
    for section in combined_sections:
        title = section.split('\n', 1)[0]  # Extract the title of the section
        if title not in seen_titles:
            seen_titles.add(title)
            unique_sections.append(section)

    # Concatenate the reordered sections
    reordered_sections = "\n".join(unique_sections)
    
    # Replace the original Directed Student Learning section with the reordered one
    start_idx = text.find(r'\subsection{Directed Student Learning}')
    end_idx = text.find(r'\subsection', start_idx + 1)
    if end_idx == -1:
        end_idx = len(text)
    
    return text[:start_idx] + r'\subsection{Directed Student Learning}' + '\n' + reordered_sections + text[end_idx:]
def reorder_student_sections2(text):
    # Define the desired order with both variations for "Master's Thesis"
    order = ["Ph.D. Dissertation", "Master's Thesis", "Master\\textquotesingle s Thesis", "Postdoctoral Mentorship", "Undergraduate Honors\nThesis"]
    
    # Extract each subsubsection based on the order
    sections = [extract_subsubsection(text, title) for title in order]
    
    # Check if any section is missing and remove it
    sections = [section for section in sections if section is not None]

    # Remove duplicate "Master's Thesis" sections (if both variations are present)
    master_thesis_sections = [section for section in sections if "Master's Thesis" in section or "Master\\textquotesingle s Thesis" in section]
    if len(master_thesis_sections) > 1:
        sections = [section for section in sections if section not in master_thesis_sections]
        sections.insert(1, master_thesis_sections[0])  # Keep only the first occurrence

    # Concatenate the reordered sections
    reordered_sections = "\n".join(sections)
    
    # Replace the original Directed Student Learning section with the reordered one
    start_idx = text.find(r'\subsection{Directed Student Learning}')
    end_idx = text.find(r'\subsection', start_idx + 1)
    if end_idx == -1:
        end_idx = len(text)
    
    return text[:start_idx] + r'\subsection{Directed Student Learning}' + '\n' + reordered_sections + text[end_idx:]




def reorder_publications(content):
    # Extract the content from \subsection{Publications} up to \subsection{Presentations}
    pattern = r'(\\subsection\{Publications\}.*?)(\\subsection\{Presentations\})'
    match = re.search(pattern, content, re.DOTALL)
    
    if not match:
        return content  # No change if the pattern isn't found
    
    publications_section = match.group(1)
    presentations_label = match.group(2)
    
    # Split into individual subsubsections
    subsubsections = re.findall(r'(\\subsubsection\{.*?\}.*?)(?=\\subsubsection|$)', publications_section, re.DOTALL)
    
    # Define the desired order
    order = ["Journal Article", "Conference Proceeding", "Book Chapters", "Other"]
    
    # Create a dictionary with subsubsection titles as keys and their content as values
    subsubsection_dict = {}
    for subsubsection in subsubsections:
        title = re.search(r'\\subsubsection\{(.*?)\}', subsubsection).group(1)
        subsubsection_dict[title] = subsubsection

    # Construct the reordered content based on the desired order
    sorted_content = ''.join([subsubsection_dict[title] for title in order if title in subsubsection_dict])
    
    # Replace old subsubsections with sorted ones in the publications section
    updated_publications = publications_section.replace(''.join(subsubsections), sorted_content)
   
    # Construct the updated content
    return content.replace(publications_section, updated_publications)

    
def capitalize_subsections(text_content):
    # Find all \subsection{...} patterns, including those spanning multiple lines
    subsections = re.findall(r'\\subsection\{(.*?)\}', text_content, re.DOTALL)

    # For each found subsection, replace it with its capitalized version
    for subsection in subsections:
        # Remove newline and excess spaces
        subsection_cleaned = ' '.join(subsection.split())
        capitalized_subsection = subsection_cleaned.upper()
        # Replace the original with the cleaned and capitalized version
        text_content = re.sub(
            r'\\subsection\{' + re.escape(subsection) + r'\}',
            r'\\subsection{' + capitalized_subsection + '}',
            text_content,
            flags=re.DOTALL
        )

    return text_content

def add_date_to_header(text_content):
    # Get the current month and year
    now = datetime.datetime.now()
    month_year = now.strftime("%B %Y")

    # Define the header settings with a darker gray color
    header_settings = r"""
\usepackage{fancyhdr}
\usepackage{xcolor}
\definecolor{darkgray}{gray}{0.4} % Define a darker gray color
\fancypagestyle{firstpage}{
    \fancyhf{} % Clear all headers and footers first
    \rhead{\textcolor{darkgray}{""" + month_year + r"""}}
    \renewcommand{\headrulewidth}{0pt} % No header rule
}
\thispagestyle{firstpage}
"""

    # Insert the header settings right before \begin{document}
    marker = r"\begin{document}"
    preamble_end = text_content.find(marker)
    
    if preamble_end != -1:
        modified_content = text_content[:preamble_end] + header_settings + text_content[preamble_end:]
        return modified_content
    else:
        print("Warning: Marker for preamble end (\\begin{document}) not found in the text content.")
        return text_content  # Return the original content if marker is not found

def add_custom_package(text, package_name="mystyle"):
    # Find the position of \author{}
    insert_pos = text.find(r'\author{')
    
    # If \author{} is found, insert the \usepackage command before it
    if insert_pos != -1:
        text = text[:insert_pos] + f"\\usepackage{{{package_name}}}\n" + text[insert_pos:]
    
    return text
def format_header(text_content):
    new_header = r"""
\begin{center}
\LARGE \textbf{\textsc{REUBEN H. KRAFT}} \\
\rule{\linewidth}{2pt}
\end{center}
\normalsize % Return to the default font size
"""
    insertion_point = text_content.find(r'\begin{document}') + len(r'\begin{document}')
    if insertion_point != -1:
        text_content = text_content[:insertion_point] + '\n' + new_header + '\n' + text_content[insertion_point:]
    return text_content
def add_date_to_header(text_content):
    now = datetime.datetime.now()
    month_year = now.strftime("%B %Y")
    header_settings = r"""
\usepackage{fancyhdr}
\usepackage{xcolor}
\definecolor{darkgray}{gray}{0.4} % Define a darker gray color
\fancypagestyle{firstpage}{
    \fancyhf{} % Clear all headers and footers first
    \rhead{\textcolor{darkgray}{""" + month_year + r"""}}
    \renewcommand{\headrulewidth}{0pt} % No header rule
}
\thispagestyle{firstpage}
"""
    insertion_point = text_content.find(r'\begin{document}')
    if insertion_point != -1:
        text_content = text_content[:insertion_point] + header_settings + text_content[insertion_point:]
    return text_content

def process_courses(text):
    course_descriptions = {
        "330": "Computational Tools for Engineers (ME 330)",
        "360": "Machine Design (ME 360)",
        "461": "Introduction to Finite Element Analysis (ME 461)",
        "563": "Nonlinear Finite Element Analysis (ME 563)",
        "497": "Development course for Computational Tools for Engineers (ME 497)",
        "440": "Capstone Design (ME 440)",
    }
    excluded_courses = {"600", "596", "496", "494", "610"}

    # Find the start and end of the course listings section
    start_of_section = text.find(r'\subsection{Teaching Experience}\label{teaching-experience}')
    end_of_section = text.find(r'\subsection{Service}\label{service}')

    # If the start or end markers are not found, return the original text
    if start_of_section == -1 or end_of_section == -1:
        print("Warning: Course section markers not found.")
        return text

    # Extract the course listings section
    course_section = text[start_of_section:end_of_section]

    # Dictionary to store courses and their years
    courses_years = {}

    # Find each year and its associated courses
    years_courses = re.findall(r'\\subsubsection\{(\d{4})\}(.*?)\\subsubsection|\\subsection', course_section, re.DOTALL)
    for year, courses in years_courses:
        for course_num in re.findall(r'(?:ME|M\s?E) (\d{3})', courses):
            if course_num not in excluded_courses:
                description = course_descriptions.get(course_num, f"ME {course_num}")
                if description not in courses_years:
                    courses_years[description] = []
                courses_years[description].append(year)

    # Format the new course section
    new_course_section = '\n\n' + r'\subsection{Teaching Experience}\label{teaching-experience}' + '\n'
    for course, years in courses_years.items():
        sorted_years = sorted(set(years), reverse=True)  # Sort and remove duplicates
        new_course_section += f"{course}, {', '.join(sorted_years)}.\n\n"
    
    # Replace the original course listings section with the new one
    new_text = text[:start_of_section] + new_course_section + text[end_of_section:]

    # Ensure a newline after \end{enumerate}
    new_text = new_text.replace('\end{enumerate}\subsection', '\end{enumerate}\n\n\subsection')
   
    return new_text

def replace_straight_quotes_with_latex_quotes(text_data):
    # This pattern matches phrases enclosed in straight double quotes, accounting for potential newlines
    pattern = r'"([\s\S]*?)"'
    
    # This function replaces straight double quotes with LaTeX quotes
    def replace_quotes(match):
        content = match.group(1)
        # Check if the content already starts with LaTeX-style quotes
        if content.startswith("``"):
            return f'"{content}"'  # Return the content as-is
        return f"``{content}''"
    
    # Use re.sub to replace all occurrences in the text
    corrected_text = re.sub(pattern, replace_quotes, text_data)
    
    # Debugging for the Service subsection
    service_section_pattern = r"\\subsection{Service}.*?\\subsubsection{Department}"
    service_section_match = re.search(service_section_pattern, corrected_text, re.DOTALL)
    
    return corrected_text

def clean_service_section(text_data):
    # Define subsubsection labels
    labels = ["College", "Department", "University", "Profession", "Society"]
    
    for label in labels:
        # Extract the subsubsection content
        section_pattern = rf"\\subsubsection{{{label}}}.*?(?=\\subsubsection|$)"
        section_match = re.search(section_pattern, text_data, re.DOTALL)
        
        if not section_match:
            print(f"{label} section not found in clean_service_section!")
            continue
        
        section_content = section_match.group(0)
        
        # Remove the redundant mention of the subsubsection title from each item
        cleaned_section = re.sub(rf"{label}, ", "", section_content)
        
        # Replace the old section with the cleaned section in the text data
        text_data = text_data.replace(section_content, cleaned_section)
    
    return text_data

def highlight_mentored_authors_astericks2(text_data):
    # Extract last names of mentored students and postdocs
    student_types = {
        'Ph.D. Dissertation': 'Ph.D. students',
        'Master\'s Thesis': 'Master students',
        'Postdoctoral Mentorship': 'Postdoc researchers',
        'Undergraduate Honors\nThesis': 'Undergrad students'
    }
    
    mentored_lastnames = []

    for subsection, label in student_types.items():
        start_index = text_data.find(rf'\subsubsection{{{subsection}}}')
        if start_index == -1:
            match = re.search(re.escape(subsection), text_data)
            if match:
                start_index = match.start()
            else:
                continue

        end_index = text_data.find(r'\subsubsection{', start_index + 1)
        if end_index == -1:
            end_index = len(text_data)

        section = text_data[start_index:end_index]
        pattern = r"\\item\s+(.*?),\s+``"
        matches = re.findall(pattern, section, re.DOTALL)
        
        for full_name in matches:
            lastname = full_name.split()[-1]
            mentored_lastnames.append(lastname)

    # Search through each publication and look for these last names
    start_publications = text_data.find(r'\subsection{Publications}')
    end_publications = text_data.find(r'\subsection{', start_publications + 1)
    if end_publications == -1:
        end_publications = len(text_data)

    publications_section = text_data[start_publications:end_publications]

    # Add a LaTeX asterisk after the name if it's found
    for lastname in mentored_lastnames:
        # Modified pattern to account for an optional middle initial
        pattern = rf"({lastname}(?:, [A-Z]\. ?[A-Z]?)?)"
        replacement = r"\1\\textsuperscript{*}"
        publications_section = re.sub(pattern, replacement, publications_section)

    # Replace the old publications section with the modified one
    modified_text = text_data[:start_publications] + publications_section + text_data[end_publications:]

    # Add mentored student note
    journal_article_section_start = modified_text.find(r'\subsubsection{Journal Article}\label{journal-article}')
    if journal_article_section_start == -1:
        print("Journal Article section not found in the text.")
        return modified_text

    note_text = "\n\\textit{Mentored student and postdoc co-authors have an astericks after their name.}\n"
    modified_text = modified_text[:journal_article_section_start] + note_text + modified_text[journal_article_section_start:]

    return modified_text

def underline_mentored_authors_with_note(text_data):
    # Extract last names of mentored students and postdocs
    student_types = {
        'Ph.D. Dissertation': 'Ph.D. students',
        'Master\'s Thesis': 'Master students',
        'Postdoctoral Mentorship': 'Postdoc researchers',
        'Undergraduate Honors\nThesis': 'Undergrad students'
    }
    
    mentored_lastnames = []

    for subsection, label in student_types.items():
        start_index = text_data.find(rf'\subsubsection{{{subsection}}}')
        if start_index == -1:
            match = re.search(re.escape(subsection), text_data)
            if match:
                start_index = match.start()
            else:
                continue

        end_index = text_data.find(r'\subsubsection{', start_index + 1)
        if end_index == -1:
            end_index = len(text_data)

        section = text_data[start_index:end_index]
        pattern = r"\\item\s+(.*?),\s+``"
        matches = re.findall(pattern, section, re.DOTALL)
        
        for full_name in matches:
            lastname = full_name.split()[-1]
            mentored_lastnames.append(lastname)

    # Search through each publication and look for these last names
    start_publications = text_data.find(r'\subsection{Publications}')
    end_publications = text_data.find(r'\subsection{', start_publications + 1)
    if end_publications == -1:
        end_publications = len(text_data)

    publications_section = text_data[start_publications:end_publications]

    # Add a LaTeX note about underlining mentored authors
    note_text = "\n\\textit{Mentored student and postdoc co-authors are underlined.}\n"
    journal_article_section_start = publications_section.find(r'\subsubsection{Journal Article}\label{journal-article}')
    publications_section = publications_section[:journal_article_section_start] + note_text + publications_section[journal_article_section_start:]

    # Underline the name if it's found
    for lastname in mentored_lastnames:
        # Modified pattern to account for an optional middle initial
        pattern = rf"({lastname}(?:, [A-Z]\. ?[A-Z]?)?)"
        replacement = r"\\underline{\1}"
        publications_section = re.sub(pattern, replacement, publications_section)

    # Replace the old publications section with the modified one
    modified_text = text_data[:start_publications] + publications_section + text_data[end_publications:]

    return modified_text

def create_custom_titles_for_sections(text_data):
    
    # change the title of the Jounral Article section
    text_data = text_data.replace(r'\subsubsection{Journal Article}\label{journal-article}', r'\subsubsection{Journal Articles}\label{journal-article}')
    
    # change the title of the Conference Proceeding section
    text_data = text_data.replace(r'\subsubsection{Conference Proceeding}\label{conference-proceeding}', r'\subsubsection{Conference Proceedings}\label{conference-proceeding}')
    
    # change the title of the Book Chapter section
    text_data = text_data.replace(r'\subsubsection{Book Chapter}\label{book-chapter}', r'\subsubsection{Book Chapters}\label{book-chapter}')

    # change the title of the Other section to Preprints and Technical Reports
    text_data = text_data.replace(r'\subsubsection{Other}\label{other}', r'\subsubsection{Preprints and Technical Reports}\label{other}')

    # in the Presentations section change the title of the Invited Talks section to Invited Talks and Seminars
    text_data = text_data.replace(r'\subsubsection{Invited}\label{invited}', r'\subsubsection{Invited Talks and Seminars}\label{invited}')

    # in the Presentations sections change the title of Uncategoried Presentations to Confererences and Workshops Presentations
    text_data = text_data.replace(r'\subsubsection{Uncategorized}\label{uncategorized}', r'\subsubsection{Conferences and Workshops}\label{uncategorized-presentations}')

    return text_data

def replace_section_colors(text, subsection_color='black', subsubsection_color='black'):
    # Define a pattern to search for the \titleformat command for \subsection and replace the color
    subsection_pattern = re.compile(r"(\\titleformat\{\\subsection\}.*?\\color\{)(blue)(\}.*)", re.DOTALL | re.VERBOSE)
    subsubsection_pattern = re.compile(r"(\\titleformat\{\\subsubsection\}.*?\\color\{)(red)(\}.*)", re.DOTALL | re.VERBOSE)

    # Define the replacement function for \subsection
    def subsection_replacement(match):
        #print("Subsection color found:", match.group(2))
        return f"{match.group(1)}{subsection_color}{match.group(3)}"

    # Define the replacement function for \subsubsection
    def subsubsection_replacement(match):
        #print("Subsubsection color found:", match.group(2))
        return f"{match.group(1)}{subsubsection_color}{match.group(3)}"

    # Perform the replacement for \subsection
    updated_text = subsection_pattern.sub(subsection_replacement, text)
    # Perform the replacement for \subsubsection
    updated_text = subsubsection_pattern.sub(subsubsection_replacement, updated_text)

    return updated_text

def process_professional_section(text):
    # Find the start of the "Professional Positions" section
    start_index = text.find("\\subsection{Professional Positions}")
    if start_index == -1:
        return "Error: Could not find the 'Professional Positions' section"
    
    # Find the end of the "Professional Positions" section
    end_index = text.find("\\subsection", start_index + 1)
    end_index = end_index if end_index != -1 else None
    
    # Extract the content before, within, and after the "Professional Positions" section
    before_professional_section = text[:start_index]
    professional_section = text[start_index:end_index]
    after_professional_section = text[end_index:] if end_index else ""
    
    # Normalize whitespace in the "Professional Positions" section
    professional_section = ' '.join(professional_section.split())

    # Example positions and locations
    positions_locations = [
        "Associate Professor of Biomedical Engineering (Courtesy), The Pennsylvania State University.",
        "Associate Professor of Mechanical Engineering, The Pennsylvania State University.",
        "Assistant Professor of Biomedical Engineering (Courtesy), The Pennsylvania State University.",
        "Assistant Professor of Mechanical Engineering, The Pennsylvania State University.",
        "Mechanical Engineer, The U.S. Army Research Laboratory, Soldier Protection Sciences Branch.",
        "Post-Doc, Oak Ridge Associated Universities at The U.S. Army Research Laboratory, Impact Physics Branch.",
        "Founder and Chief Engineer, BrainSim Technologies Inc.",
        "Lead Researcher of Computational Biomechanics, The Johns Hopkins University Applied Physics Laboratory, Research and Exploratory Development Department, Biomechanics and Injury Mitigation Systems Group."
    ]

    # Process each position and location
    for pl in positions_locations:
        normalized_pl = ' '.join(pl.split())
        
        # Find the position and location in the content
        position_start = professional_section.find(normalized_pl)
        if position_start == -1:
            print("Error: Could not find position and location:", pl)
            continue
        
        # Extract the position, location, and date info
        position_end = position_start + len(normalized_pl)
        rest_of_content = professional_section[position_end:]
        date_match = re.search(r'\((.*?)\)', rest_of_content)
        
        if not date_match:
            print("Error: Could not find date information for:", pl)
            continue
        
        date_info = date_match.group(1)
        updated_date_info = process_dates(date_info)
        
        # Replace old date string with new date format
        date_start = position_end + rest_of_content.find(date_info)
        date_end = date_start + len(date_info) + 2  # +2 to include parentheses
        updated_entry = normalized_pl + " (" + updated_date_info + "). "
        updated_entry += "\n"


        professional_section = professional_section[:position_start] + updated_entry + professional_section[date_end:]
        
    # Reconstruct the entire text string with the updated "Professional Positions" section
    updated_text = before_professional_section + professional_section + after_professional_section
    
    # Add line returns for better readability in the .tex source file
    updated_text = updated_text.replace('. ', '. \n')

    # Add line returns after specific subsections and subsubsections
    updated_text = updated_text.replace("\\subsection{Professional Positions}\label{professional-positions}", "\\subsection{Professional Positions}\label{professional-positions}\n")
    updated_text = updated_text.replace("\\subsubsection{Academic}\label{academic}", "\\subsubsection{Academic}\label{academic}\n")


    return updated_text


def process_dates(date_string):
    # Define possible date formats
    date_formats = ["%B %Y", "%B %d, %Y"]

    # Define the function for processing dates
    def process_single_date(date):
        for fmt in date_formats:
            try:
                return datetime.datetime.strptime(date, fmt).strftime("%Y")
            except ValueError:
                continue
        return "Present" if "Present" in date else "Invalid Date"

    date_ranges = date_string.split(" - ")
    start_date = process_single_date(date_ranges[0])
    end_date = process_single_date(date_ranges[1]) if len(date_ranges) > 1 else "Present"

    return f"{start_date} - {end_date}"
def format_positions(text):
    # Define the pattern to match the positions and dates
    pattern = re.compile(r'(.*?), The (.*?)\.\s+\((.*?)\)\.')

    def replace_with_tabular(match):
        # Extract the position, location, and date
        position = match.group(1).strip()
        location = "The " + match.group(2).strip()
        date = match.group(3).strip()

        # Return the formatted tabular environment string
        return (
            r"\begin{tabular}{@{}l@{\hskip 0.5in}r@{}}"
            f"{position}, {location}. & ({date}). \\\\"
            r"\end{tabular}"
        )

    # Search and replace all occurrences of the pattern with the tabular format
    updated_text = re.sub(pattern, replace_with_tabular, text)

    return updated_text


def format_professional_positions_section(text, spacing='5pt'):
    # Add the booktabs package if not present
    if '\\usepackage{booktabs}' not in text:
        text = text.replace(r'\usepackage{xcolor}', r'\usepackage{xcolor}' + '\n' + r'\usepackage{booktabs}')

    # Insert the tabularx package right after the xcolor package if not present
    if '\\usepackage{tabularx}' not in text:
        text = text.replace(r'\usepackage{xcolor}', r'\usepackage{xcolor}' + '\n' + r'\usepackage{tabularx}')

    # Pattern to extract the Professional Positions section
    positions_section_pattern = r'(\\subsection\{Professional Positions\}\\label\{professional-positions\})(.+?)(\\subsection\{Education\}\\label\{education\})'
    match = re.search(positions_section_pattern, text, re.DOTALL)

    if match:
        # Extract the subsection title and the 'Education' section header
        subsection_title = match.group(1).strip()
        education_section_header = match.group(3).strip()

        # Initialize an empty string to store the formatted Professional Positions section
        formatted_positions = subsection_title + "\n\n"

        professional_positions_section = match.group(2)

        # Split the section into subsubsections, capturing the titles
        subsubsections = re.split(r'(\\subsubsection\{.*?\}\\label\{.*?\})', professional_positions_section)

        # Process each subsubsection and its entries
        for i in range(1, len(subsubsections), 2):
            subsubsection_title = subsubsections[i].strip()
            formatted_positions += subsubsection_title + "\n\n"

            # Begin the tabularx environment for the subsubsection
            formatted_positions += "\\begin{tabularx}{\\textwidth}{@{}Xr@{}}\n"

            # Split into entries
            entries = subsubsections[i + 1].strip().split('\n\n')
            for index, entry in enumerate(entries):
                entry = entry.strip()
                if entry:
                    # Use regular expression to separate the position/location from the date range
                    date_range_pattern = r'\(\d{4} - (?:\d{4}|Present)\)'
                    date_range_match = re.search(date_range_pattern, entry)
                    if date_range_match:
                        date_range = date_range_match.group(0).strip('()')  # Remove parentheses
                        pos_loc = entry[:date_range_match.start()].strip()
                        formatted_positions += f"{pos_loc} & {date_range}"
                        # Add the line break and space unless it's the last entry
                        if index < len(entries) - 1:
                            formatted_positions += f" \\\\\n\\addlinespace[{spacing}]\n"
                    else:
                        # Handle the case where the date range is not found
                        formatted_positions += f"{entry} & N/A"
                        # Add the line break and space unless it's the last entry
                        if index < len(entries) - 1:
                            formatted_positions += f" \\\\\n\\addlinespace[{spacing}]\n"

            # End the tabularx environment for the subsubsection
            formatted_positions += "\\end{tabularx}\n\n"

        # Replace the original Professional Positions section with the formatted section
        text = text.replace(match.group(0), formatted_positions + education_section_header)

    else:
        print("The 'Professional Positions' subsection was not found.")

    return text

def format_education_section(text, spacing='5pt'):
    # Define the pattern to find the education section
    education_section_pattern = r'(\\subsection\{Education\}\\label\{education\})(.+?)(\\subsection\{.*?\}\\label\{.*?\}|\Z)'
    match = re.search(education_section_pattern, text, re.DOTALL)
   
    if match:
        subsection_title = match.group(1).strip()
        following_section_header = match.group(3).strip()
        education_entries = match.group(2).strip().split('\n\n')

        # Initialize an empty string to store the formatted Education section
        formatted_education = subsection_title + "\n\n\\begin{flushleft}\n\\begin{tabularx}{\\textwidth}{@{}Xr@{}}\n"
        
        # Process each entry
        for index, entry in enumerate(education_entries):
            if entry:
                # Remove any trailing period from each entry before the last line break
                entry = re.sub(r'\.\s*(?=\\|$)', '', entry)
                # Split entry into institution/major and year
                parts = entry.rsplit(',', 1)
                if len(parts) == 2:
                    institution_major = parts[0].strip()
                    year = parts[1].strip().lstrip('\\').rstrip('.')
                    formatted_education += f"{institution_major} & {year}"
                else:
                    formatted_education += f"{entry} & N/A"

                # Add the line break and space unless it's the last entry
                if index < len(education_entries) - 1:
                    formatted_education += f" \\\\\n\\addlinespace[{spacing}]\n"

        # Close the tabularx and flushleft environments
        formatted_education += "\n\\end{tabularx}\n\\end{flushleft}\n\n" + following_section_header
        # print(formatted_education)
        # Replace the original Education section with the formatted section
        text = text.replace(match.group(0), formatted_education)

    else:
        print("The 'Education' subsection was not found.")

    return text

def format_awards_and_honors_section(text, spacing='5pt', borders=False):
    # Add the longtable package if not present
    if '\\usepackage{longtable}' not in text:
        text = text.replace(r'\usepackage{tabularx}', r'\usepackage{tabularx}' + '\n' + r'\usepackage{longtable}')

    # Pattern to extract the Awards and Honors section
    awards_section_pattern = r'(\\subsection\{Awards and Honors\}\\label\{awards-and-honors\})(.+?)(\\subsection\{Publications\}\\label\{publications\})'
    match = re.search(awards_section_pattern, text, re.DOTALL)

    if match:
        # Extract the subsection title and the 'Publications' section header
        subsection_title = match.group(1).strip()
        publications_section_header = match.group(3).strip()

        # Initialize an empty string to store the formatted Awards and Honors section
        formatted_awards = subsection_title + "\n\n"

        # Define new lengths for table columns
        #formatted_awards += "\\newlength{\\mylength}\n"
        #formatted_awards += "\\newlength{\\myotherlength}\n"
        #formatted_awards += "\\setlength{\\mylength}{\\dimexpr(\\textwidth-0\\tabcolsep)*3/4\\relax}\n"
        #formatted_awards += "\\setlength{\\myotherlength}{\\dimexpr(\\textwidth-2\\tabcolsep)/4\\relax}\n"
        
        # Define the fixed width for the second column
        fixed_second_column_width = "1.0in"  # or the width you prefer

        # Modify the formatted_awards string with the new lengths
        #formatted_awards = ""
        formatted_awards += "\\newlength{\\mylength}\n"
        formatted_awards += "\\newlength{\\myotherlength}\n"
        formatted_awards += f"\\setlength{{\\myotherlength}}{{{fixed_second_column_width}}}\n"  # Set the fixed width for the second column
        formatted_awards += "\\setlength{\\mylength}{\\dimexpr(\\textwidth-\\myotherlength-0\\tabcolsep)\\relax}\n"  # Calculate the remaining width for the first column

        # Define the table format string based on whether borders are enabled
        if borders:
            table_format = "|p{\\mylength}|>{\\raggedleft\\arraybackslash}p{\\myotherlength}|"
        else:
            table_format = "@{}p{\\mylength}@{}>{\\raggedleft\\arraybackslash}p{\\myotherlength}@{}"

        # Begin the longtable environment for the section
        formatted_awards += f"\\begin{{longtable}}{{{table_format}}}\n"
        if borders:
            formatted_awards += "\\hline\n"

        awards_and_honors_section = match.group(2).strip()

        # Split into entries by looking for two consecutive newlines
        entries = awards_and_honors_section.strip().split('\n\n')
        for index, entry in enumerate(entries):
            entry = entry.strip()
            if entry:
                # Use regular expression to separate the award/honor description from the date
                date_pattern = r'\((?:\w+ )?(\d{4})(?: - (\d{4}|\w+))?\.?\)'
                entry = re.sub(r'\s+', ' ', entry)  # Replace multiple whitespaces with a single space
                date_match = re.search(date_pattern, entry)
                if date_match:
                    # Extract the year or year range
                    year = date_match.group(1)
                    award_desc = entry[:date_match.start()].strip()
                    # Add vertical space before each entry except the first one
                    if index > 0:
                        formatted_awards += "\n\\addlinespace[{}]\n".format(spacing)
                    formatted_awards += f"{award_desc} & {year} \\\\"
                    if borders and index < len(entries) - 1:
                        formatted_awards += "\\hline\n"
                else:
                    # Handle the case where the date is not found
                    if index > 0:
                        formatted_awards += "\n\\addlinespace[{}]\n".format(spacing)
                    formatted_awards += f"{entry} & N/A \\\\"
                    if borders and index < len(entries) - 1:
                        formatted_awards += "\\hline\n"

        # End the longtable environment for the section
        formatted_awards += "\\end{longtable}\n"

        # Replace the original Awards and Honors section with the formatted section
        text = text.replace(match.group(0), formatted_awards + publications_section_header)

    else:
        print("The 'Awards and Honors' subsection was not found.")

    return text

def update_documentclass_font_size(latex_content, font_size):
    # Define the pattern to find the documentclass declaration
    pattern = re.compile(r'\\documentclass\[(.*?)\]\{article\}', re.DOTALL)
    
    # Define the replacement string to include the desired font size option
    replacement = fr'\\documentclass[{font_size}pt]{{article}}'
    
    # Replace the documentclass declaration with the one including the desired font size option
    updated_content = re.sub(pattern, replacement, latex_content)
    
    return updated_content

def create_template_latex_file(filename):
    with open(filename, 'w') as file:
        file.write(r"""
\documentclass[a4paper,10pt]{article}
\usepackage[a4paper, margin=1in]{geometry}
\usepackage{amsmath,amssymb}
\usepackage{iftex}

\ifPDFTeX
  \usepackage[T1]{fontenc}
  \usepackage[utf8]{inputenc}
  \usepackage{textcomp}
  \usepackage{newtxtext,newtxmath} % Times New Roman-like font for pdfLaTeX
\else
  \usepackage{fontspec} % Allows font customization
  \setmainfont{Times New Roman} % Set Times New Roman as main font for XeLaTeX/LuaLaTeX
\fi
                   
\usepackage{lmodern}
\ifPDFTeX\else
  % xetex/luatex font selection
\fi
\IfFileExists{upquote.sty}{\usepackage{upquote}}{}
\IfFileExists{microtype.sty}{%
  \usepackage[]{microtype}
  \UseMicrotypeSet[protrusion]{basicmath} % disable protrusion for tt fonts
}{}
\makeatletter
\@ifundefined{KOMAClassName}{%
  \IfFileExists{parskip.sty}{%
    \usepackage{parskip}
  }{
    \setlength{\parindent}{0pt}
    \setlength{\parskip}{6pt plus 2pt minus 1pt}}
}{
  \KOMAoptions{parskip=half}}
\makeatother
\usepackage{xcolor}
\usepackage{tabularx}
\usepackage{longtable}
\usepackage{booktabs}
\setlength{\emergencystretch}{3em}
\providecommand{\tightlist}{%
  \setlength{\itemsep}{0pt}\setlength{\parskip}{0pt}}
\setcounter{secnumdepth}{-\maxdimen}
\ifLuaTeX
  \usepackage{selnolig}
  \setmainfont{Times New Roman}
\fi
\usepackage{bookmark}
\IfFileExists{xurl.sty}{\usepackage{xurl}}{}
\urlstyle{same}
\hypersetup{
  hidelinks,
  pdfcreator={LaTeX via pandoc}}

\author{}
\date{}
\begin{document}
\end{document}
""")
        

def add_custom_package(text_content, package_name="mystyle"):
    insertion_point = text_content.find(r'\author{')
    if insertion_point != -1:
        text_content = text_content[:insertion_point] + f"\\usepackage{{{package_name}}}\n" + text_content[insertion_point:]
    return text_content

def set_section_colors(text_content):
    color_commands = r"""
    \usepackage{titlesec}
    \usepackage{color}
    \titleformat{\subsection}
    {\normalfont\large\bfseries\color{black}} % format
    {\thesubsection} % label
    {1em} % sep
    {} % before-code
    \titleformat{\subsubsection}
    {\normalfont\normalsize\bfseries\color{black}} % format
    {\thesubsubsection} % label
    {1em} % sep
    {} % before-code
    """
    insertion_point = text_content.find(r'\begin{document}')
    if insertion_point != -1:
        text_content = text_content[:insertion_point] + color_commands + text_content[insertion_point:]
    return text_content

def read_word_document(file_path):
    doc = Document(file_path)
    full_text = []
    table_data = []

    # Read paragraphs
    for para in doc.paragraphs:
        full_text.append(unidecode.unidecode(para.text))

    # Read tables
    for table in doc.tables:
        
        table_rows = []
        for row in table.rows:
            row_text = [unidecode.unidecode(cell.text.strip()) for cell in row.cells]
            
            table_rows.append(row_text)
        table_data.append(table_rows)  
    
    return '\n'.join(full_text), table_data

def extract_text_between_markers(full_text, start_marker, end_marker):
    # Escape special characters in start and end markers for regex
    start_marker = re.escape(start_marker)
    end_marker = re.escape(end_marker)
    
    # Compile the regex pattern
    pattern = re.compile(f'{start_marker}(.*?){end_marker}', re.DOTALL)
    # Search for the pattern
    match = pattern.search(full_text)
    if match:
        
        return match.group(1).strip()
    else:
        return ""

def extract_text_up_to_end_marker(word_text, start_marker, end_marker):
    """
    Extract the text starting from the `start_marker` to the `end_marker`, inclusive of the end marker.
    """
    # Find the position of the start_marker and end_marker
    start_pos = word_text.find(start_marker)
    if start_pos == -1:
        return None  # Start marker not found
    
    end_pos = word_text.find(end_marker, start_pos)
    if end_pos == -1:
        return None  # End marker not found
    
    # Extract the text from the start_marker to the end_marker (inclusive)
    return word_text[start_pos:end_pos + len(end_marker)]


def process_courses_from_word(word_text, latex_text):
    # Define markers for extraction
    start_marker = "List of Credit Courses Taught at Penn State for Each Semester with Enrollments in Each Course"
    end_marker = "Concise Compilation of Results of Student Feedback from Multiple Sources"

    # Extract the relevant courses section
    relevant_text = extract_text_between_markers(word_text, start_marker, end_marker)

    # Define course descriptions and excluded courses
    course_descriptions = {
        "330": "Computational Tools for Engineers",
        "360": "Machine Design",
        "461": "Introduction to Finite Element Analysis",
        "563": "Nonlinear Finite Element Analysis",
        "497": "Development course for Computational Tools for Engineers",
        "440": "Capstone Design",
    }
    excluded_courses = {"600", "596", "496", "494", "610"}

    # Dictionary to store years and list of courses for each year
    years_courses = defaultdict(set)  # Use a set to automatically handle duplicates

    current_year = None
    for line in relevant_text.split('\n'):
        # Match the year (Spring, Summer, Fall Year)
        if re.match(r'(Spring|Summer|Fall) \d{4}', line):
            current_year = line.split()[-1]
        # Match the course number (ME XXX), allowing spaces between M and E
        elif re.match(r'M\s*E\s*\d{3}', line):
            match = re.search(r'M\s*E\s*(\d{3})', line)
            course_num = match.group(1)
            if course_num not in excluded_courses:
                description = course_descriptions.get(course_num, f"ME {course_num}")
                if current_year:
                    years_courses[current_year].add((course_num, description))  # Store tuples of (course_num, description)
        # For courses without a code (like "Development course for Computational Tools for Engineers")
        else:
            if current_year and "course" in line.lower():
                years_courses[current_year].add((None, line.strip()))  # Add the full name without 

    # Sort the years in descending order
    sorted_years = sorted(years_courses.keys(), reverse=True)

    # Format the new course section with bold year and courses in new line
    new_course_section = '\\subsection*{TEACHING EXPERIENCE}\\label{teaching-experience}\n'
    for year in sorted_years:
        # Add the year in bold with a new line after
        new_course_section += f"\\textbf{{{year}}}\\\\\n"  # Add explicit line break (\\)
        # Sort courses by course number in descending order, place None at the end
        sorted_courses = sorted(years_courses[year], key=lambda x: (x[0] is None, x[0]), reverse=True)
        # Format each course (code first, then name in brackets if there's a code)
        formatted_courses = []
        for course_num, description in sorted_courses:
            if course_num:
                formatted_courses.append(f"ME {course_num} ({description})")
            else:
                formatted_courses.append(f"{description}")
        # Add the courses, separated by commas
        new_course_section += ', '.join(formatted_courses) + '\n\n'

    # Insert the new course section before \end{document}
    new_text = latex_text.replace("\\end{document}", new_course_section + "\\end{document}")
    
    return new_text

# def process_courses_from_word( word_text,latex_text):
   

#     # Define markers
#     start_marker = "List of Credit Courses Taught at Penn State for Each Semester with Enrollments in Each Course"
#     end_marker = "Concise Compilation of Results of Student Feedback from Multiple Sources, Documented Evaluation of Candidate’s Programs, Activities, and Skills in Relating to Clientele"

#     # Extract the relevant courses section
#     relevant_text = extract_text_between_markers(word_text, start_marker, end_marker)

#     # Define course descriptions and excluded courses
#     course_descriptions = {
#         "330": "Computational Tools for Engineers (ME 330)",
#         "360": "Machine Design (ME 360)",
#         "461": "Introduction to Finite Element Analysis (ME 461)",
#         "563": "Nonlinear Finite Element Analysis (ME 563)",
#         "497": "Development course for Computational Tools for Engineers (ME 497)",
#         "440": "Capstone Design (ME 440)",
#     }
#     excluded_courses = {"600", "596", "496", "494", "610"}

#     # Dictionary to store courses and their years
#     courses_years = {}

#     current_year = None
#     for line in relevant_text.split('\n'):
#         if re.match(r'(Spring|Summer|Fall) \d{4}', line):
#             current_year = line.split()[-1]
#         elif re.match(r'ME\s?\d{3}', line):
#             match = re.search(r'ME\s?(\d{3})', line)
#             course_num = match.group(1)
#             if course_num not in excluded_courses:
#                 description = course_descriptions.get(course_num, f"ME {course_num}")
#                 if description not in courses_years:
#                     courses_years[description] = []
#                 if current_year and current_year not in courses_years[description]:
#                     courses_years[description].append(current_year)

#     # Format the new course section
#     new_course_section = '\\subsection{TEACHING EXPERIENCE}\\label{teaching-experience}\n'
#     for course, years in courses_years.items():
#         sorted_years = sorted(set(years), reverse=True)  # Sort and remove duplicates
#         new_course_section += f"{course}, {', '.join(sorted_years)}.\n\n"

#     # Insert the new course section before \end{document}
#     new_text = latex_text.replace("\\end{document}", new_course_section + "\\end{document}")
   
#     return new_text

def format_publication_entry(publication):
    # Extract the DOI and format it as a URL
    
    doi_match = re.search(r"DOI: (\S+)", publication)
    if doi_match:
        doi_url = f"\\url{{https://doi.org/{doi_match.group(1)}}}"
        publication = re.sub(r"DOI: \S+", "", publication)
    else:
        doi_url = ""
    # Adding emphasis to the author name "Kraft, R. H."
    publication = re.sub(r'(Kraft,)( R\.\s*H\.)', r'\\textbf{\\textbf{\1}\2}', publication)
    

    # Combine the formatted publication entry
    formatted_entry = f"\\item {publication.strip()} {doi_url}".strip()
    formatted_entry = formatted_entry.replace('&', '&\n')
    return formatted_entry

def replace_special_characters(text):
    # Mapping of accented characters to LaTeX-friendly equivalents
    replacements = {
        'á': r"\'a",
        'é': r"\'e",
        'í': r"\'i",
        'ó': r"\'o",
        'ú': r"\'u",
        'ñ': r"\~n",
        'ü': r"\"u",
        'Á': r"\'A",
        'É': r"\'E",
        'Í': r"\'I",
        'Ó': r"\'O",
        'Ú': r"\'U",
        'Ñ': r"\~N",
        'Ü': r"\"U",
        'ç': r"\c{c}",
        'Ç': r"\c{C}",
        'ö': r"\"o",
        'Ö': r"\"O",
        # Add other replacements as needed
        '’': "'",
    }
    
    # Replace characters in the text
    for char, latex_equiv in replacements.items():
        text = text.replace(char, latex_equiv)
    
    return text


import re

import re

def extract_publications(word_text, latex_text):
    start_marker_journal = "Journal Article"
    end_marker_journal = "Parts of Books"
    start_marker_conference = "Conference Proceedings"
    end_marker_conference = "Other Works"
    start_marker_book = "Book Chapter"
    end_marker_book = "Refereed Conference Proceedings"
    start_marker_other = "Other Works"
    end_marker_other = "Manuscripts Submitted for Publication"

    journal_publications = extract_text_between_markers(word_text, start_marker_journal, end_marker_journal)
    conference_publications = extract_text_between_markers(word_text, start_marker_conference, end_marker_conference)
    book_chapters = extract_text_between_markers(word_text, start_marker_book, end_marker_book)
    other_publications = extract_text_between_markers(word_text, start_marker_other, end_marker_other)
    
    def underline_students(text):
        # Regular expression to capture names with "Graduate Student" or "Undergraduate Student" designation
        student_pattern = r'([A-Z][a-zA-Z\s\.,]+) \((?:Primary Author|Co-Author|Student Author) - (Graduate Student|Undergraduate Student|Postdoctoral Student)\)'
        # Replace matched patterns with LaTeX underline command
        return re.sub(student_pattern, r'\\underline{\1}', text)

    latex_output = r"""
\subsection{PUBLICATIONS}\label{publications}

\subsubsection{Journal Article}\label{journal-article}

\begin{enumerate}
\def\labelenumi{\arabic{enumi}.}
"""

    for publication in journal_publications.split("\n"):
        if publication.strip():
            # Clean up and format publication entry
            publication = replace_special_characters(format_publication_entry(publication))
            publication = underline_students(publication)  # Apply underline to student names
            formatted_publication = publication.replace("&", r"\&")

            # Remove numbers after \item and add to LaTeX output
            formatted_publication = re.sub(r'^\\item\s*\d+\.', r'\\item', formatted_publication)
            latex_output += f"  {formatted_publication}\n"

    latex_output += r"""
\end{enumerate}

\subsubsection{Conference Proceeding}\label{conference-proceeding}

\begin{enumerate}
\def\labelenumi{\arabic{enumi}.}
"""
    for publication in conference_publications.split("\n"):
        if publication.strip():
            # Clean up and format publication entry
            publication = replace_special_characters(format_publication_entry(publication))
            publication = underline_students(publication)  # Apply underline to student names
            formatted_publication = publication.replace("&", r"\&")

            # Remove numbers after \item and add to LaTeX output
            formatted_publication = re.sub(r'^\\item\s*\d+\.', r'\\item', formatted_publication)
            latex_output += f"  {formatted_publication}\n"

    latex_output += r"""
\end{enumerate}

\subsubsection{Book Chapters}\label{book-chapters}

\begin{enumerate}
\def\labelenumi{\arabic{enumi}.}
"""
    for publication in book_chapters.split("\n"):
        if publication.strip():
            # Clean up and format publication entry
            publication = replace_special_characters(format_publication_entry(publication))
            publication = underline_students(publication)  # Apply underline to student names
            formatted_publication = publication.replace("&", r"\&")

            # Remove numbers after \item and add to LaTeX output
            formatted_publication = re.sub(r'^\\item\s*\d+\.', r'\\item', formatted_publication)
            latex_output += f"  {formatted_publication}\n"

    latex_output += r"""
\end{enumerate}

\subsubsection{Other}\label{other}

\begin{enumerate}
\def\labelenumi{\arabic{enumi}.}
"""
    for publication in other_publications.split("\n"):
        if publication.strip() and not publication.startswith(r"Pre-Print") and not publication.startswith(r"Technical Report"):
            # Clean up and format publication entry
            publication = replace_special_characters(format_publication_entry(publication))
            publication = underline_students(publication)  # Apply underline to student names
            formatted_publication = publication.replace("&", r"\&")

            # Remove numbers after \item and add to LaTeX output
            formatted_publication = re.sub(r'^\\item\s*\d+\.', r'\\item', formatted_publication)
            latex_output += f"  {formatted_publication}\n"

    latex_output += r"""
\end{enumerate}
"""
    
    # Append the new publication section to the existing LaTeX content
    latex_text = latex_text + latex_output

    # Remove the existing \end{document} if it exists
    latex_text = latex_text.replace("\\end{document}", "").strip()

    latex_text += r"\vspace{1\baselineskip}"

    # Append the end of the LaTeX document
    latex_text += "\n\\end{document}"

    return latex_text




def clean_input_text(text):
    """Remove unsupported characters from the input text."""
    return text.encode('utf-8', 'ignore').decode('utf-8')

def replace_special_characters(entry):
    replacements = {
        'á': r"\'a",
        'é': r"\'e",
        'í': r"\'i",
        'ó': r"\'o",
        'ú': r"\'u",
        'ñ': r"~n",
        'ü': r"\"u",
        'Ä': r"\"A",
        'Ö': r"\"O",
        'Ü': r"\"U",
        'ä': r"\"a",
        'ö': r"\"o",
        'ß': r"\ss",
        # Add more replacements as needed
    }
    for key, value in replacements.items():
        entry = entry.replace(key, value)
    return entry



def extract_presentations(word_text, latex_text):
    # Define the markers for each section
    section_markers = {
        'Demonstrations': ('Demonstrations', 'Keynotes/Plenary Addresses'),
        'Keynotes/Plenary Addresses': ('Keynotes/Plenary Addresses', 'Oral Presentations'),
        'Oral Presentations': ('Oral Presentations', 'Panels'),
        'Panels': ('Panels', 'Posters'),
        'Posters': ('Posters', 'Posters and Oral Presentations'),
        'Posters and Oral Presentations': ('Posters and Oral Presentations', 'Seminars'),
        'Seminars': ('Seminars', "Description of Outreach or Other Activities in which there was Significant Use of Candidate's Expertise"),
    }

    # Dictionary to store unique entries
    unique_entries = set()

    # Begin LaTeX output for presentations
    latex_output = r"""
\subsection{Presentations}\label{presentations}

\begin{enumerate}
\def\labelenumi{\arabic{enumi}.}
"""

    # Helper function to determine if a line is a label or an actual entry
    def is_content_line(line):
        # Ignore lines that are just section names or irrelevant headings
        headings = {"Oral Presentations", "Panels", "Posters", "Seminars", "Posters and Oral Presentations", "and Workshops",'Invited Keynote'}
        return bool(line.strip()) and line.strip() not in headings and "Won Best Poster Award" and ", and Workshops" not in line

    # Extract text between markers for each section and add unique entries only
    for section, (start_marker, end_marker) in section_markers.items():
        extracted_text = extract_text_between_markers(word_text, start_marker, end_marker)
        
        for entry in extracted_text.splitlines():
            entry = entry.strip()
            if is_content_line(entry) and entry not in unique_entries:  # Ensure entry is unique and skip irrelevant lines
                unique_entries.add(entry)  # Track it to prevent future duplicates
                formatted_entry = replace_special_characters(format_publication_entry(entry))
                formatted_entry = formatted_entry.replace("&", r"\&")
                formatted_entry = re.sub(r'^\\item\s*\d+\.', r'\\item', formatted_entry)  # Clean item numbering
                latex_output += f"  {formatted_entry}\n"

    latex_output += r"""
\end{enumerate}
"""

    # Insert the new presentation section into the LaTeX text
    latex_text += latex_output

    # Remove existing \end{document} if it exists and add it at the end
    latex_text = latex_text.replace("\\end{document}", "").strip()
    latex_text += r"\vspace{1\baselineskip}\n\\end{document}"

    return latex_text

def extract_contract_project_and_grants(word_text, latex_text):
    # Define markers for extracting text
    start_marker_grants = 'Projects, Grants, Commissions, and Contracts'
    end_marker_grants = 'Submitted for Funding: September 27, 2024'
    extracted_text = extract_text_between_markers(word_text, start_marker_grants, end_marker_grants)
    
    # Filter out unwanted lines and ensure entries are unique
    unwanted_lines = {
        'Pending', 
        'Agency: Triad National Security, LLC (was LANL - Los Alamos National Laboratory)',
        'Principal Investigator: Kraft, Reuben H.',
        'Co-Investigator(s):',
        'Project Title: Elucidating high strain rate deformation mechanisms in penetration-resistant composites.',
        'Won Best Poster Award'
    }
    filtered_text = "\n".join(
        line for line in extracted_text.splitlines() if line.strip() and line not in unwanted_lines
    )

    # Prepare the LaTeX formatted output for the required section
    latex_output = r"""
                \subsection{CONTRACT, FELLOWSHIPS, GRANTS AND SPONSORED RESEARCH}\label{contract-research}
                """
    
    # Track unique entries in the grants section
    current_agency = ""
    seen_lines = set()
    
    for line in filtered_text.splitlines():
        line = line.strip()
        if line.startswith("Agency:"):
            if current_agency and line not in seen_lines:
                latex_output += "\\\\[12pt]\n"  # Add space before new agency section
            
            current_agency = line
            latex_output += f"{current_agency}\\\\\n"
            seen_lines.add(line)
        
        elif line.startswith("Principal Investigator:") or line.startswith("Co-Investigator(s):") or line.startswith("Project Title:") or line.startswith("Amendments:") or "OSP Number" in line:
            if line not in seen_lines:
                latex_output += f"{line}\\\\\n"
                seen_lines.add(line)

    # Finalize LaTeX document
    latex_text = latex_text.replace("\\end{document}", "").strip()
    latex_text += latex_output
    latex_text += "\n\\end{document}"

    return latex_text







def process_student_thesis_titles(text,dossier_file_path): 
    # Extract master's thesis titles
    masters_data = extract_student_titles(dossier_file_path)
 
    # Extract PhD dissertation titles
    phd_data = extract_phd_titles(dossier_file_path)
  
    # Extract postdoc titles
    postdoc_data = extract_postdoc_titles(dossier_file_path)
    
    # Extract undergrad thesis titles
    undergrad_data = extract_undergrad_student_titles(dossier_file_path)
     
    # Merge the dictionaries
    student_data = {**masters_data, **phd_data, **postdoc_data, **undergrad_data}

    # Replace problematic characters in titles
    student_data = replace_problematic_characters_in_titles(student_data)
        
    # Update the LaTeX content with thesis titles - does not work for postdoc titles see next function
    text = add_title_to_name2(text, student_data)
       
    # reformat phd section for name, title, date
    text = reformat_phd_section(text)
    
    # reformat masters section for name, title, date    
    text = update_masters_section(text)
    
    # reformat postdoc section for name, title, date
    text = update_postdoc_section(text)
    
    # reformat undergrad section for name, title, date
    text = update_undergraduate_section(text)
        
    return text

def replace_problematic_characters_in_titles(student_data):
    problematic_char = u'\u2013'  # En dash (U+2013)

    for name, title in student_data.items():
        if problematic_char in title:
            #print(f"Found character '{problematic_char}' in {name}: {title}")
            student_data[name] = title.replace(problematic_char, '--')

    return student_data

def update_undergraduate_section(text_data):
    # Define section label
    label = "Undergraduate Honors\nThesis"
    
    # Extract the Undergraduate Honors Thesis section
    section_pattern = rf"\\subsubsection{{{label}}}.*?\\begin{{enumerate}}.*?\\end{{enumerate}}"
    section_match = re.search(section_pattern, text_data, re.DOTALL)
    
    if not section_match:
        print("Undergraduate Honors Thesis section not found!")
        return text_data
    
    section_content = section_match.group(0)
    
    # Extract individual items
    item_pattern = r"Undergraduate Honors Thesis\. \((.*?)\).\\\\\s+Advised: (.*?), \"(.*?)\","
    items = re.findall(item_pattern, section_content, re.DOTALL)
    
    # Format items
    formatted_items = []
    for time_period, name, title in items:
        formatted_item = f"\\item {name}, \"{title}\", {time_period}"
        formatted_items.append(formatted_item)
    
    # Combine formatted items
    formatted_section = "\\subsubsection{Undergraduate Honors\nThesis}\n\\begin{enumerate}\n\\def\\labelenumi{\\arabic{enumi}.}\n" + "\n".join(formatted_items) + "\n\\end{enumerate}"
    
    # Replace the old section with the new formatted section in the text data
    updated_text_data = text_data.replace(section_content, formatted_section)
    
    return updated_text_data

def update_postdoc_section(text_data):
    # Define section label
    label = "Postdoctoral Mentorship"
    
    # Extract the Postdoctoral Mentorship section
    section_pattern = rf"\\subsubsection{{{label}}}.*?\\begin{{enumerate}}.*?\\end{{enumerate}}"
    section_match = re.search(section_pattern, text_data, re.DOTALL)
    
    if not section_match:
        print("Postdoctoral Mentorship section not found!")
        return text_data
    
    section_content = section_match.group(0)
    
    # Extract individual items
    item_pattern = r"Postdoctoral Mentorship\. \((.*?)\).\\\\\s+Advised: (.*?), \"(.*?)\","
    items = re.findall(item_pattern, section_content, re.DOTALL)
    
    # Format items
    formatted_items = []
    for time_period, name, title in items:
        formatted_item = f"\\item {name}, \"{title}\", {time_period}"
        formatted_items.append(formatted_item)
    
    # Combine formatted items
    formatted_section = "\\subsubsection{Postdoctoral Mentorship}\n\\begin{enumerate}\n\\def\\labelenumi{\\arabic{enumi}.}\n" + "\n".join(formatted_items) + "\n\\end{enumerate}"
    
    # Replace the old section with the new formatted section in the text data
    updated_text_data = text_data.replace(section_content, formatted_section)
    
    return updated_text_data

def update_masters_section(text_data):
    # Define the two possible formats for "Master's Thesis"
    master_thesis_options = ["Master's Thesis", "Master\\textquotesingle s Thesis"]

    # Detect which format is used in the text
    master_thesis_format = next((option for option in master_thesis_options if option in text_data), None)

    if master_thesis_format is None:
        print("Master's Thesis section format not detected!")
        return text_data

    # Escape backslashes for regex
    escaped_master_thesis_format = master_thesis_format.replace('\\', '\\\\')

    # Prepare regex pattern for the section
    section_pattern = rf"(\\subsubsection\{{.*?{escaped_master_thesis_format}.*?\}}.*?\\begin\{{enumerate\}})(.*?)(\\end\{{enumerate\}})"
    
    section_match = re.search(section_pattern, text_data, re.DOTALL)

    if not section_match:
        print("Master's Thesis section not found!")
        return text_data

    section_start, section_content, section_end = section_match.groups()
    
    # Adjusting item pattern for matching items within the section
    item_pattern = rf"\\item\s*{escaped_master_thesis_format}\.\s*\((.*?)\)\.\\\\\s*Advised:\s*(.*?),\s*\"(.*?)\""
    
    items = re.findall(item_pattern, section_content, re.DOTALL | re.MULTILINE)
    
    if not items:
        print("No items found in Master's Thesis section!")
        return text_data

    # Format items
    formatted_items = [f"\\item {name}, \"{title}\", {time_period}" for time_period, name, title in items]

    # Combine formatted items
    formatted_section_content = "\n".join(formatted_items)

    # Replace the old section with the new formatted section in the text data
    updated_section = section_start + formatted_section_content + section_end
    updated_text_data = text_data.replace(section_content, formatted_section_content)
    
    return updated_text_data
def extract_student_titles(file_path):
    #file_path = r"C:\Users\rhk12\OneDrive - The Pennsylvania State University\resume\CV\20231021-095357-CDT.docx"
    
    # Check if the file exists
    if not os.path.exists(file_path):
        return f"File not found at '{file_path}'"
    
    # If the file exists, proceed to read its content
    doc = Document(file_path)
    full_text = [para.text for para in doc.paragraphs]

    # Define the sub-section titles
    sub_section_titles = ["Master's Thesis Advisor", "Master\u2019s Thesis Advisor"]
    end_section_titles = ["Master's Thesis Committee Member", "Master\u2019s Thesis Committee Member"]

    # Find the start of the sub-section
    start_index = None
    for title in sub_section_titles:
        start_index = next((i for i, text in enumerate(full_text) if title in text), None)
        if start_index is not None:
            break

    if start_index is None:
        print(f"Section 'Master's Thesis Advisor' not found")
        return

    # Find the end of the sub-section
    end_index = None
    for title in end_section_titles:
        end_index = next((i for i, text in enumerate(full_text[start_index:]) if title in text), None)
        if end_index is not None:
            break

    if end_index is None:
        end_index = len(full_text)
    else:
        end_index += start_index

    # Extract student names and thesis titles
    student_data = {}
    for line in full_text[start_index + 1:end_index]:
        if ", " in line and "MS." in line:
            name = line.split(",")[0].strip()
            title_start = line.find("MS.") + 4
            title_end = line.find(".", title_start)
            title = line[title_start:title_end].strip()
            student_data[name] = title
    
    return student_data
def add_undergrad_titles(text, undergrad_data):
    for name, title in undergrad_data.items():
        # Convert the name format from 'Last, F.' to 'First Last'
        last_name, first_initial = name.split(', ')
        # We'll use a placeholder for the first name since we don't have the full first name
        first_name_placeholder = f"{first_initial}.*"
        
        # Define the pattern to search for the student's entry in the LaTeX content
        pattern = rf"(Advised: {first_name_placeholder} {last_name})"
        
        # Check if the pattern exists in the text
        if re.search(pattern, text):
            # Replace the pattern with the pattern + title
            text =  re.sub(pattern, f"\\1, \"{title}\"", text)
            print(f"Updated {first_initial}. {last_name}'s entry.")
        else:
            print(f"Pattern for {first_initial}. {last_name} not found in the text!")
    
    return text

def extract_postdoc_titles(file_path):
    #file_path = r"C:\Users\rhk12\OneDrive - The Pennsylvania State University\resume\CV\20231021-095357-CDT.docx"
    
    # Check if the file exists
    if not os.path.exists(file_path):
        return f"File not found at '{file_path}'"
    
    # If the file exists, proceed to read its content
    doc = Document(file_path)
    full_text = [para.text for para in doc.paragraphs]

    # Define the sub-section titles
    sub_section_title = "Postdoctoral Mentorship Advisor"
    end_section_title = "Research Activity Advisor"

    # Find the start of the sub-section
    start_index = next((i for i, text in enumerate(full_text) if sub_section_title in text), None)
    if start_index is None:
        return {}

    # Find the end of the sub-section
    end_index = next((i for i, text in enumerate(full_text[start_index:]) if end_section_title in text), None)
    if end_index is None:
        end_index = len(full_text)
    else:
        end_index += start_index

    # Extract postdoc names and work titles
    postdoc_data = {}
    for line in full_text[start_index + 1:end_index]:
        if ". " in line:
            name, title_with_dates = line.split(". ", 1)
            title = title_with_dates.split(". (")[0].strip()
            postdoc_data[name.strip()] = title

    return postdoc_data

def extract_phd_titles(file_path):
    #file_path = r"C:\Users\rhk12\OneDrive - The Pennsylvania State University\resume\CV\20231021-095357-CDT.docx"
    
    # Check if the file exists
    if not os.path.exists(file_path):
        return f"File not found at '{file_path}'"
    
    # If the file exists, proceed to read its content
    doc = Document(file_path)
    full_text = [para.text for para in doc.paragraphs]

    # Define the sub-section titles
    sub_section_title = "Ph.D. Dissertation Advisor"
    end_section_title = "Ph.D. Dissertation Committee Member"  # This is the section that follows the Advisor section

    # Find the start of the sub-section
    start_index = next((i for i, text in enumerate(full_text) if sub_section_title in text), None)
    if start_index is None:
        print(f"Section '{sub_section_title}' not found")
        return {}

    # Find the end of the sub-section
    end_index = next((i for i, text in enumerate(full_text[start_index + 1:]) if end_section_title in text), None)
    if end_index is None:
        end_index = len(full_text)
    else:
        end_index += start_index + 1

    # Extract student names and dissertation titles
    student_data = {}
    for line in full_text[start_index + 1:end_index]:
        if ", " in line and "Ph.D." in line:
            name = line.split(",")[0].strip()
            title_start = line.find("Ph.D.") + 6
            title_end = line.find(".", title_start)
            title = line[title_start:title_end].strip()
            student_data[name] = title

    return student_data

def convert_table_to_latex(table):
        latex_table = "\\begin{tabular}{|" + " | ".join(["c"] * len(table[0])) + "|}\n\\hline\n"
        for row in table:
            latex_table += " & ".join(row) + " \\\\\n\\hline\n"
        latex_table += "\\end{tabular}"
        return latex_table

import re

def add_professional_positions_to_latex(latex_text, tables_list, header_to_search):
    def format_professional_positions(table):
        formatted_positions = {
            "Academic": [],
            "Government": [],
            "Professional": []
        }
        
        for row in table[1:]:  # Skip the header row
            if len(row) < 4:
                continue  # Skip rows that don't have enough columns

            employer = row[0].replace("\n", " ").strip()
            rank_or_title = row[2].strip()
            dates = row[3].strip()
            
            # Extract months and years from the dates
            date_range = re.findall(r'(?:\w+\s)?\d{4}', dates)
            if len(date_range) == 2:
                formatted_dates = f"({date_range[0]} - {date_range[1]})"
            elif len(date_range) == 1:
                formatted_dates = f"({date_range[0]})"
            else:
                formatted_dates = ""

            # Construct the full position line with the desired format
            position_line = f"{rank_or_title}, {employer}. {formatted_dates}. \\\\ \n"

            # Classify positions based on keywords
            if "Assistant Professor" in rank_or_title or "Professor" in rank_or_title:
                formatted_positions["Academic"].append(position_line)
            elif "U.S. Army Research Laboratory" in employer:
                formatted_positions["Government"].append(position_line)
            else:
                formatted_positions["Professional"].append(position_line)

        return formatted_positions

    def extract_current_position(table):
        # Checks if the table has a header that matches the "Exact Rank and Title" table
        if len(table) > 1 and "Exact Rank and Title of Position" in table[0]:
            for row in table:
                if len(row) >= 3 and "Professor of Mechanical Engineering" in row[2]:
                    formatted_title = "Professor of Mechanical Engineering. (July 2024 - Present)."
                    return formatted_title
        return None

    # Remove \end{document} if it exists
    latex_text = latex_text.replace("\\end{document}", "").strip()
    flag = True
    for table in tables_list:
        # Extract and add the current position if available
        if flag:
            exact_rank_title = extract_current_position(table)
            if exact_rank_title:
                flag = False
        
        # Check for the professional positions table and process it
        if table and len(table) > 1 and table[0][0] == "Previous Employers with City/State\nIncluding U.S. Military\n(Most Recent First)":
            latex_text += "\n\\subsection*{PROFESSIONAL POSITIONS}\n"
            
            # Convert and append the formatted professional positions
            positions = format_professional_positions(table)
            
            # Add Academic subsubsection
            if positions["Academic"]:
                latex_text += "\\subsubsection*{Academic}\n"
                
                # Insert current position under Academic if it exists
                if exact_rank_title:
                    latex_text += f"\\noindent {exact_rank_title}\\\\\n\\vspace{{-0.5\\baselineskip}}\n\n"
                
                for entry in positions["Academic"]:
                    latex_text += "\\noindent\\begin{tabular*}{\\textwidth}{@{\\extracolsep{\\fill}} p{0.85\\textwidth} l} \n"
                    latex_text += entry 
                    latex_text += "\\end{tabular*} \\vspace{-0.5\\baselineskip}\n\n"
            
            # Add Government subsubsection
            if positions["Government"]:
                latex_text += "\\subsubsection*{Government}\n"
                for entry in positions["Government"]:
                    latex_text += "\\noindent\\begin{tabular*}{\\textwidth}{@{\\extracolsep{\\fill}} p{0.85\\textwidth} l} \n"
                    latex_text += entry 
                    latex_text += "\\end{tabular*} \\vspace{-0.5\\baselineskip}\n\n"
            
            # Add Professional subsubsection
            if positions["Professional"]:
                latex_text += "\\subsubsection*{Professional}\n"
                for entry in positions["Professional"]:
                    latex_text += "\\noindent\\begin{tabular*}{\\textwidth}{@{\\extracolsep{\\fill}} p{0.85\\textwidth} l} \n"
                    latex_text += entry 
                    latex_text += "\\end{tabular*} \\vspace{-0.5\\baselineskip}\n\n"

    # Append end of LaTeX document
    latex_text += "\n\\end{document}"

    return latex_text







def read_word_document(file_path):
    doc = Document(file_path)
    full_text = []
    table_data = []

    # Read paragraphs
    for para in doc.paragraphs:
        full_text.append(unidecode.unidecode(para.text))

    # Read tables
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            row_text = [unidecode.unidecode(cell.text.strip()) for cell in row.cells]
            table_rows.append(row_text)
        table_data.append(table_rows)  # Store the entire table
    
    return '\n'.join(full_text), table_data

# #here is the function to extract proffession_position
# def add_professional_positions_to_latex(latex_text, tables_list, header_to_search):
#     def format_professional_positions(table):
#         formatted_positions = {
#             "Academic": [],
#             "Government": [],
#             "Professional": []
#         }
        
#         for row in table[1:]:
#             employer = row[0].replace("\n", " ").strip()  # Remove newlines within employer text
#             work_performed = row[1].strip()               # Get the work description
#             rank_or_title = row[2].strip()                # Extract the rank or title
#             dates = row[3].strip()                        # Get the date range
            
#             # Remove months and keep only years from the date string using regex
#             date_years = re.findall(r'\d{4}', dates)
#             formatted_dates = " - ".join(date_years)      # Format as "start year - end year"

#             # Classify positions based on keywords
#             if "Assistant Professor" in rank_or_title or "Professor" in rank_or_title:
#                 formatted_positions["Academic"].append(f"{rank_or_title}, {employer} & {formatted_dates} \\\\ \n")
#             elif "U.S. Army Research Laboratory" in employer:
#                 formatted_positions["Government"].append(f"{rank_or_title}, {employer} & {formatted_dates} \\\\ \n")
#             else:
#                 formatted_positions["Professional"].append(f"{rank_or_title}, {employer} & {formatted_dates} \\\\ \n")

#         return formatted_positions
    
#     # Remove the existing \end{document} if it exists
#     latex_text = latex_text.replace("\\end{document}", "").strip()

#     # Iterate through the list of tables
#     for table in tables_list:
#         # Check if the header exists in the first row of the table
#         if header_to_search in table[0]:
#             # Add a new section to the LaTeX document for "Professional Positions"
#             latex_text += "\n\\subsection*{PROFESSIONAL POSITIONS}\n"
            
#             # Convert and append the formatted professional positions
#             positions = format_professional_positions(table)
            
#             # Add Academic subsubsection
#             if positions["Academic"]:
#                 latex_text += "\\subsubsection*{Academic}\n"
#                 for entry in positions["Academic"]:
#                     latex_text += "\\noindent\\begin{tabular*}{\\textwidth}{@{\\extracolsep{\\fill}} p{0.85\\textwidth} l} \n"
#                     latex_text += entry 
#                     if entry != positions["Academic"][-1]:
#                         latex_text += "\\end{tabular*} " + "\\vspace{-0.4\\baselineskip} \n\n"
#                     else:
#                         latex_text += "\\end{tabular*} " + "\\vspace{-0.5\\baselineskip} \n\n"
            
#             # Add Government subsubsection
#             if positions["Government"]:
                
#                 latex_text += "\\subsubsection*{Government}\n"
#                 for entry in positions["Government"]:
#                     latex_text += "\\noindent\\begin{tabular*}{\\textwidth}{@{\\extracolsep{\\fill}} p{0.85\\textwidth} l} \n"
#                     latex_text += entry 
#                     if entry != positions["Government"][-1]:
#                         latex_text += "\\end{tabular*} " + "\\vspace{-0.4\\baselineskip} \n\n"
#                     else:
#                         latex_text += "\\end{tabular*} " + "\\vspace{-0.5\\baselineskip} \n\n"
            
#             # Add Professional subsubsection
#             if positions["Professional"]:
                
#                 latex_text += "\\subsubsection*{Professional}\n"
#                 for entry in positions["Professional"]:
#                     latex_text += "\\noindent\\begin{tabular*}{\\textwidth}{@{\\extracolsep{\\fill}} p{0.85\\textwidth} l} \n"
#                     latex_text += entry 
#                     if entry != positions["Professional"][-1]:
#                         latex_text += "\\end{tabular*} " + "\\vspace{-0.4\\baselineskip} \n\n"
#                     else:
#                         latex_text += "\\end{tabular*} " + "\\vspace{-0.5\\baselineskip} \n\n"

#     # Append the end of the LaTeX document
#     latex_text += "\n\\end{document}"
    
#     return latex_text


def add_education_section(latex_text, tables, column_name):
    # Check for the header in each table
    for table in tables:
        # Get the headers from the first row of the table
        headers = table[0]
        
        # Check if the column name exists in the headers
        if column_name in headers:
            # Remove the existing \end{document} if it exists
            latex_text = latex_text.replace("\\end{document}", "").strip()
            
            # Find the index of the institution, major, and date columns
            institution_index = headers.index('Name and City/State of Institution')
            major_index = headers.index('Major Subjects')
            date_index = headers.index('Degrees - Dates')
            
            # Start creating the LaTeX section
            education_section = '\n\n\\section*{EDUCATION}\n'
            
            # Iterate over the rows (skip the header)
            for row in table[1:]:
                # Replace "Post-Doctoral" with "Concentration" in all relevant columns
                degree = row[date_index].split(',')[0].strip().replace("Post-Doctoral", "Concentration")
                institution = row[institution_index].strip()
                major = row[major_index].strip().replace("Post-Doctoral", "Concentration")
                dates = row[date_index].split(',')[-1].strip()  # Extracting the date part

                # Extract months and years from the dates
                date_range = re.findall(r'\b(?:\w+\s)?\d{4}\b', dates)  # Capture full year range
                formatted_date = ""
                if len(date_range) == 2:
                    formatted_date = f"{date_range[0]} - {date_range[1]}"
                elif len(date_range) == 1:
                    formatted_date = f"{date_range[0]}"

                # Format: Degree, Institution, Location, Date in one line, Major in the next
                education_section += f'{degree}, {institution}, {formatted_date} \\\\ \n'
                education_section += f'Major: {major} \\\\[0.2cm]\n'
            
            # Add the education section to the LaTeX text
            latex_text += education_section
            
            # Append the end of the LaTeX document only if it's not already there
            if "\\end{document}" not in latex_text:
                latex_text += "\n\\end{document}"

    return latex_text


def format_award_entry(entry):
    # Remove unnecessary newlines and extract the first sentence (main award title and date)
    entry_lines = entry.split("\n")
    first_line = entry_lines[0].strip()  # Only take the first line for each award

    # Regex to match different date formats
    year_match = re.search(r'\((\d{4})\)', first_line)  # Match single year in parentheses (e.g., (2019))
    
    if not year_match:
        range_match = re.search(r'\((\d{4})\s*-\s*(\d{4})\)', first_line)  # Match year range (e.g., (2013 - 2016))
        if range_match:
            year_match = range_match.group(2)  # Extract only the ending year of the range
        else:
            month_year_match = re.search(r'\(\w+\s+(\d{4})\)', first_line)  # Match month and year (e.g., (August 2023))
            if month_year_match:
                year_match = month_year_match.group(1)  # Extract only the year part (2023)

    year = year_match if isinstance(year_match, str) else year_match.group(1) if year_match else ""  # Get the matched year or ending year
    first_line = re.sub(r'\((\w+\s+\d{4})\)\.|\((\d{4})\)\.|\((\d{4})\s*-\s*(\d{4})\)\.', '', first_line)  # Remove the date or range from the main text
    
    return (first_line.strip(), year)

def extract_awards(document_text):
    # Split document text into individual awards based on double newlines (paragraphs)
    awards = document_text.strip().split("\n\n")
    
    formatted_awards = []
    for award in awards:
        # Format each award entry
        formatted_award, award_year = format_award_entry(award)
        formatted_awards.append(f"{formatted_award} & {award_year} \\\\\n")
    
    return formatted_awards

def add_awards_and_honors(latex_text, document_text):
    # Remove the existing \end{document} if it exists
    latex_text = latex_text.replace("\\end{document}", "").strip()

    # Initialize the awards and honors section
    awards_section = "\\subsection*{AWARDS AND HONORS}\n\n"

    # Define markers to extract relevant sections from the document
    markers = [
        ("Honors or Awards for Excellence in Teaching and Advising\n\nTeaching",
         "Supervision of, and Membership on,"),
        ("Honors or Awards for Scholarship, Research, or Creative Activities\n\nScholarship/Research",
         "Technology Transferred or Adapted in the Field"),
        ("Honors or Awards for Leadership and/or Service to the University, Community, or the Profession\n\nLeadership",
         "Service, Professional"),
        ("Service, Professional\n\n",
         "EXTERNAL LETTERS OF ASSESSMENT"),
    ]

    # Extract and combine all relevant text
    relevant_text = ""
    for start_marker, end_marker in markers:
        relevant_text += extract_text_between_markers(document_text, start_marker, end_marker) + "\n\n"
        relevant_text.lstrip(',').strip()
    # Split relevant text into individual awards
    awards = [award.strip() for award in relevant_text.strip().split("\n\n") if award.strip()]

    # Iterate over awards and format each entry
    for award in awards:
        # Clean up double periods and ensure consistent formatting
        award = award.replace("..", ".").strip()
        award = award.lstrip(',').strip()

        # Add a line break after the date using regex
        # Handles both year and month-year patterns, inserting a LaTeX line break
        award = re.sub(r"(\(\d{4}(?: - \d{4})?\))(\.?)(?!\n)", r"\1.\n\\", award)  # Handles (2018), (2013 - 2016)
        award = re.sub(r"(\(\w+ \d{4}\))(\.?)(?!\n)", r"\1.\n\\", award)  # Handles (August 2023), (October 2023)

        # Add a line break between different awards
        awards_section += f"{award}"

    # Remove excessive line breaks and extra spaces
    awards_section = "\n\n".join(line for line in awards_section.splitlines() if line.strip())

    # Append the awards section to the LaTeX text
    latex_text += "\n" + awards_section

    # Append the end of the LaTeX document
    latex_text += "\n\n\\end{document}"

    return latex_text


def format_award_entry(award_text):
    # Extract the date from the end of the entry
    match = re.search(r'\(([^)]+)\)$', award_text)
    if match:
        award_date = match.group(1).strip()
        award_title = award_text[:match.start()].strip()
    else:
        award_date = ""
        award_title = award_text.strip()
    return award_title, award_date




def add_impact_in_society(word_document, latex_text):
    start_marker = 'Impact in Society of Research Scholarship and Creative Accomplishment'
    end_marker = 'Record of Membership in Professional and Learned Societies'
    
    # Extract the text between the specified markers
    text = extract_text_between_markers(word_document, start_marker, end_marker)
    
    # Clean and format the extracted text for LaTeX
    if text:
        latex_text = latex_text.replace("\\end{document}", "").strip()
        formatted_text = text.replace("&", r"\&").replace("%", r"\%").replace("#", r"\#")  # Escape LaTeX special characters
        formatted_text = formatted_text.strip()  # Remove any leading/trailing whitespace

        # Add a new subsection to the LaTeX document
        latex_text += r"""
\subsection*{Impact in Society of Research, Scholarship, and Creative Accomplishment}
\noindent
""" + formatted_text + "\n\\vspace{0.5\\baselineskip}\n"
        
    latex_text += "\n\\end{document}"

    # Return the updated LaTeX document text
    return latex_text


def add_service_to_uni(word_text, latex_text):
    start_marker = 'Record of Committee Work at Department, Division, School, Campus, College, and University Levels'
    end_marker = 'Service to Society as a Representative of the University'
    text = extract_text_between_markers(word_text, start_marker, end_marker)

    if text:
        # Prepare LaTeX formatted text
        latex_formatted_text = latex_text.replace("\\end{document}", "").strip()
        latex_formatted_text += "\n\\section*{Service to the University}\n\n"
        
        # Define headings and their subheadings
        headings = {
            'College': [
                'Academic Leadership and Support Work',
                'Committee Work',
                'Competition Judging'
            ],
            'Department': [
                'Academic Leadership and Support Work',
                'Committee Work'
            ],
            'University': [
                'Committee Work',
                'Participation in Development/Fundraising Activities'
            ]
        }

        # Split the text into lines
        lines = text.strip().splitlines()
        current_heading = None
        current_subheading = None
        added_headings = set()  # Track which headings have been added

        for line in lines:
            line = line.strip()
            
            # Check for main headings and add them only once
            if line in headings and line not in added_headings:
                current_heading = line
                added_headings.add(line)
                latex_formatted_text += f"\\subsection*{{{line}}}\n\n"
                current_subheading = None  # Reset current subheading
                
            # Check for subheading content based on keywords
            elif 'Committee' in line and current_heading:
                if current_subheading != 'Committee Work':  # Add the Committee Work subheading only once under the current heading
                    current_subheading = 'Committee Work'
                    latex_formatted_text += "\\hspace{1cm}\\subsubsection*{Committee Work}\n\n"
                latex_formatted_text += line + "\n\n"
            elif 'Research' in line and current_heading == 'Department':
                if current_subheading != 'Academic Leadership and Support Work':  # Add Academic Leadership and Support Work under Department
                    current_subheading = 'Academic Leadership and Support Work'
                    latex_formatted_text += "\\hspace{1cm}\\subsubsection*{Academic Leadership and Support Work}\n\n"
                latex_formatted_text += line + "\n\n"
            elif 'Development' in line and current_heading == 'University':
                if current_subheading != 'Participation in Development/Fundraising Activities':
                    current_subheading = 'Participation in Development/Fundraising Activities'
                    latex_formatted_text += "\\hspace{1cm}\\subsubsection*{Participation in Development/Fundraising Activities}\n\n"
                latex_formatted_text += line + "\n\n"
            elif line:  # Any non-empty line not caught by previous conditions
                latex_formatted_text += line + "\n\n"

        # Add the formatted text to the LaTeX text and end the document
        latex_text = latex_formatted_text + "\n\\end{document}"
        
    return latex_text

####################################################################################################


def add_service_to_society_as_rep_of_uni(word_text, latex_text):
    # Define specific markers for extraction
    start_marker = 'Judged Posters for Undergraduate Exhibition\n\nService to Society as a Representative of the University'
    end_marker = 'Service to the Disciplines and to the Profession'
    
    # Extract the text between the markers
    text = extract_text_between_markers(word_text, start_marker, end_marker)

    if text:
        # Prepare LaTeX formatted text
        latex_formatted_text = latex_text.replace("\\end{document}", "").strip()
        latex_formatted_text += "\n\\section*{Service to Society as a Representative of the University}\n\n"
        
        # Split text into lines for processing
        lines = text.strip().splitlines()
        current_subheading = None

        for line in lines:
            line = line.strip()
            
            # Check for subheadings and add to LaTeX with smaller size
            if line == 'Participation in Community Affairs':
                current_subheading = line
                latex_formatted_text += f"\\subsubsection*{{{line}}}\n\n"
            elif line == 'Service to Governmental Agencies at the International, Federal, State, or Local Levels':
                current_subheading = line
                latex_formatted_text += f"\\subsubsection*{{{line}}}\n\n"
                
            # Format each entry under subheading or directly
            elif line:
                latex_formatted_text += "\\hspace{1cm}" + line + "\n\n"
        
        # Finalize the LaTeX text
        latex_text = latex_formatted_text + "\n\\end{document}"

    return latex_text


def add_service_to_discipline_and_to_the_profession(word_text, latex_text):
    # Define specific markers for extraction
    start_marker = 'Service to the Disciplines and to the Profession'
    end_marker = 'Co-Organizer and Co-Chairperson, 2014 Mid-Atlantic Section (M-AS) of the American Physical Society (APS), Co-Organizer. (January 2014 - October 2014).'
    
    # Extract the text between markers including the end marker
    text = extract_text_up_to_end_marker(word_text, start_marker, end_marker)

    if text:
        # Prepare LaTeX formatted text
        latex_formatted_text = latex_text.replace("\\end{document}", "").strip()
        latex_formatted_text += "\n\\section*{Service to the Disciplines and to the Profession}\n\n"
        
        # Split text into lines for processing
        lines = text.strip().splitlines()
        current_subheading = None

        for line in lines:
            line = line.strip()

            # Check for subheadings and add to LaTeX with smaller size
            if 'Organizing Conferences and Service on Conference Committees' in line:
                current_subheading = line
                latex_formatted_text += f"\\subsubsection*{{{line}}}\n\n"
            
            # Skip the duplicated section heading if it appears in the extracted content
            elif line != 'Service to the Disciplines and to the Profession' and line:  # Non-empty lines are added
                latex_formatted_text += "\\hspace{1cm}" + line + "\n\n"
        
        # Finalize the LaTeX text
        latex_text = latex_formatted_text + "\n\\end{document}"

    return latex_text

def add_directed_student_learning(word_text, latex_text):
    # Define the markers and their corresponding LaTeX subheadings
    sections = [
        ("Master's Thesis Advisor", "Master's Thesis Committee Member"),
        ("Master's Thesis Committee Member", "Ph.D. Dissertation Advisor"),
        ("Ph.D. Dissertation Advisor", "Ph.D. Dissertation Committee Member"),
        ("Ph.D. Dissertation Committee Member", "Postdoctoral Mentorship Advisor"),
        ("Postdoctoral Mentorship Advisor", "Research Activity Advisor"),
        ("Research Activity Advisor", "Undergraduate Honors Thesis Advisor"),
        ("Undergraduate Honors Thesis Advisor", 
         "Robinson, M., Undergraduate. The development of an anatomically correct model of calcaneus fracture and")
    ]
    
    # LaTeX section title
    directed_learning_section = "\\subsection*{DIRECTED STUDENT LEARNING}\n\n"
    
    # Extract and process each section
    for start_marker, end_marker in sections:
        # Extract the text between the markers
        extracted_text = extract_text_between_markers(word_text, start_marker, end_marker)
        
        # Add the LaTeX subheading
        directed_learning_section += f"\\subsubsection*{{{start_marker}}}\n"
        
        # Split the extracted text into lines
        lines = extracted_text.splitlines()
        for line in lines:
            line = line.strip()
            if not line:
                continue  # Skip empty lines
            
            # Add a double newline after lines with "Date Graduated" for better readability
            if "Date Graduated:" in line:
                directed_learning_section += f"{line}\n\n"
            else:
                # Other lines are added with a single newline for spacing
                directed_learning_section += f"{line}\n\n"
        
        # Add extra space between subsections for clarity
        directed_learning_section += "\n\n"
    
    # Append the new section to the LaTeX content
    latex_text = latex_text.replace("\\end{document}", "").strip()  # Remove existing \end{document}
    latex_text += "\n\n" + directed_learning_section + "\n\n\\end{document}"  # Add the new section and re-add \end{document}
    
    return latex_text




def main():
    #open a file to read in C:\Users\rhk12\OneDrive - The Pennsylvania State University\resume\CV named main.tex 
    # open the file for reading
    #f = open(r'C:\Users\rhk12\OneDrive - The Pennsylvania State University\resume\CV\main.tex', 'r')
    #f = open(r'C:\Users\rhk12\OneDrive - The Pennsylvania State University\resume\CV\main.tex', 'r', encoding='utf-8')
    # Read the document
    word_text, table_data = read_word_document("CV_Data.docx")

    # Write the content to a text file
    # with open("doc.txt", "w", encoding="utf-8") as file:
    #     file.write(table)

    # with open(word_text, 'r') as file:
    #     file.read()
    
    filename = 'output.tex'
    create_template_latex_file(filename)
    
    with open(filename, 'r') as file:
        text2 = file.read()
    
    #add my formating package
    text2 = add_custom_package(text2)
    
    #Set the section colors
    text2 = set_section_colors(text2)
      
     # Format the header
    text2 = format_header(text2)
    
    # add the date to the latex file
    text2 = add_date_to_header(text2)

    text2 = add_education_section(text2, table_data, "Degrees - Dates")

    text2 = add_awards_and_honors(text2, word_text)

    text2 = extract_publications(word_text,text2)

    text2 = extract_presentations(word_text, text2)

    text2 = add_professional_positions_to_latex(text2, table_data, "Professor of Mechanical Engineering")

    text2 = extract_contract_project_and_grants(word_text,text2)

    text2 = process_courses_from_word(word_text, text2)

    text2 = add_directed_student_learning(word_text,text2)

    text2 = add_impact_in_society(word_text,text2)
    
    text2 = add_service_to_uni(word_text,text2)

    text2 = add_service_to_society_as_rep_of_uni(word_text, text2)

    text2 = add_service_to_discipline_and_to_the_profession(word_text, text2)
    
    # write_courses_to_file(text, f1)
    
    # text = process_courses(text)
    
    # reorder the publications
    # text = reorder_publications(text)
    
    # # reorder the student sections
    # text = reorder_student_sections4(text)
    
    # # process the student thesis titles
    # text = process_student_thesis_titles(text,"University+Dossier-20240524-075749-CDT.docx")
    with open(filename, 'w') as file:
        file.write(text2)
    # # clean up the service section to remove extra text
    #text = clean_service_section(text)

    # # replace straight quotes with latex quotes
    # #text = replace_straight_quotes_with_latex_quotes(text)
    
    # # add emphasis for mentees  - pick between astericks or underlines
    # #text = highlight_mentored_authors_astericks2(text)
    # text = underline_mentored_authors_with_note(text)
    
    # # format the professional section   
    #text = process_professional_section(text)
    
    # # convert the professional positions section to a tabularx environment
    # # so the date range can be right justified
    # text = format_professional_positions_section(text,spacing='5pt')

    # # convert the education section to a tabularx environment
    # # so the date range can be right justified
    # text = format_education_section(text, spacing='5pt')
    
    # # format the awards and honors section
    # # so the date range can be right justified
    # text = format_awards_and_honors_section(text, spacing='5pt', borders=False)
    
    # # ------------------------------------------
    # # --------- Final Formatting ----------------
    # # ------------------------------------------

    # # add custom titles for sections
    # text = create_custom_titles_for_sections(text)
    
    # # colorize the subsection and subsubsection titles
    # text = replace_section_colors(text, subsection_color='black', subsubsection_color='black')
    
    # # make subsection text uppercase - needs to be last
    # text = capitalize_subsections(text)
    
    # text = update_documentclass_font_size(text, 11)

  
    # f1.write(text)
    
    # #close the files
    # f.close()
    # f1.close()
    

if __name__ == "__main__":
    main()

    