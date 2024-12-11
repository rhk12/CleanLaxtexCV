import re
import datetime
import datetime
from docx import Document 
import unidecode

def paragraphs_text(file_path):
    doc = Document(file_path)
    full_text = []

    for para in doc.paragraphs:
        full_text.append(unidecode.unidecode(para.text))
    
    return '\n'.join(full_text)

def get_table_data(doc, table_index):
    table_data = []
    table = doc.tables[table_index]
    keys = None

    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        is_bold_row = all(
            all(run.bold for run in cell.paragraphs[0].runs if run.text.strip())
            for cell in row.cells if cell.paragraphs
        )

        if is_bold_row:
            keys = tuple(text)
            continue

        if keys:
            row_data = dict(zip(keys, text))
            table_data.append(row_data)

    return table_data

def extract_text_between_markers(full_text, start_marker, end_marker=None):
    start_marker = re.escape(start_marker)
    if end_marker:
        end_marker = re.escape(end_marker)
        pattern = re.compile(f'{start_marker}(.*?){end_marker}', re.DOTALL)
    else:
        pattern = re.compile(f'{start_marker}(.*)', re.DOTALL)
    
    match = pattern.search(full_text)
    if match:
        return match.group(1).strip()
    else:
        return ""

def create_template_latex_file(filename):
    with open(filename, 'w') as file:
        file.write(r"""
\documentclass[a4paper,10pt]{article}
\usepackage[a4paper, margin=1in]{geometry}
\usepackage{amsmath,amssymb}
\usepackage{iftex}
\usepackage{enumitem} 

\ifPDFTeX
  \usepackage[T1]{fontenc}
  \usepackage[utf8]{inputenc}
  \usepackage{textcomp}
  \usepackage{newtxtext,newtxmath} % Times New Roman-like font for pdfLaTeX
\else
  \usepackage{fontspec} % Allows font customization
  \setmainfont{Times New Roman} % Set Times New Roman as main font for XeLaTeX/LuaLaTeX
\fi
                   
\usepackage{lmodern}
\ifPDFTeX\else
  % xetex/luatex font selection
\fi
\IfFileExists{upquote.sty}{\usepackage{upquote}}{}
\IfFileExists{microtype.sty}{%
  \usepackage[]{microtype}
  \UseMicrotypeSet[protrusion]{basicmath} % disable protrusion for tt fonts
}{}
\makeatletter
\@ifundefined{KOMAClassName}{%
  \IfFileExists{parskip.sty}{%
    \setlength{\parskip}
  }{
    \setlength{\parindent}{0pt}
    \setlength{\parskip}{6pt}}
}{
\KOMAoptions{parskip=half}}
\makeatother
\usepackage{xcolor}
\usepackage{tabularx}
\usepackage{longtable}
\usepackage{booktabs}
\setlength{\emergencystretch}{3em}
\providecommand{\tightlist}{%
  \setlength{\itemsep}{0pt}\setlength{\parskip}{0pt}}
\setcounter{secnumdepth}{-\maxdimen}
\ifLuaTeX
  \usepackage{selnolig}
  \setmainfont{Times New Roman}
\fi
\usepackage{bookmark}
\IfFileExists{xurl.sty}{\usepackage{xurl}}{}
\urlstyle{same}
\hypersetup{
  hidelinks,
  pdfcreator={LaTeX via pandoc}}

\author{}
\date{}
\begin{document}
{{positions}}
{{education}}
{{awards_and_honor}}
{{publications}}
{{awarded}}
{{Intellectual}}
{{directedstudent}}
{{service}}
{{editorialboard}}
{{professionalmembership}}
\end{document}
""")


def set_section_colors(text_content):
    color_commands = r"""
    \usepackage{titlesec}
    \usepackage{color}
    \titleformat{\subsection}
    {\normalfont\large\bfseries\color{black}} % format
    {\thesubsection} % label
    {1em} % sep
    {} % before-code
    \titleformat{\subsubsection}
    {\normalfont\normalsize\bfseries\color{black}} % format
    {\thesubsubsection} % label
    {1em} % sep
    {} % before-code
    """
    insertion_point = text_content.find(r'\begin{document}')
    if insertion_point != -1:
        text_content = text_content[:insertion_point] + color_commands + text_content[insertion_point:]
    return text_content


def format_header(text_content, doc):
    profile_table_index = 0
    table_data = get_table_data(doc, profile_table_index)
    profile_data = {}
    for entry in table_data:
        profile_data.update(entry)

    first_name = profile_data.get('First Name and Initial', '')
    last_name = profile_data.get('Last Name', '')

    full_name = first_name + last_name

    new_header = rf"""
    \begin{{center}}
    \LARGE \textbf{{\textsc{{{full_name.upper()}}}}} \\
    \rule{{\linewidth}}{{2pt}}
    \end{{center}}
    \normalsize % Return to the default font size
    """

    insertion_point = text_content.find(r'\begin{document}') + len(r'\begin{document}')
    if insertion_point != -1:
        text_content = text_content[:insertion_point] + '\n' + new_header + '\n' + text_content[insertion_point:]
    
    return text_content


def add_custom_package(text_content, package_name="mystyle"):
    insertion_point = text_content.find(r'\author{')
    if insertion_point != -1:
        text_content = text_content[:insertion_point] + f"\\usepackage{{{package_name}}}\n" + text_content[insertion_point:]
    return text_content


def add_date_to_header(text_content):
    now = datetime.datetime.now()
    month_year = now.strftime("%B %Y")
    header_settings = r"""
    \usepackage{fancyhdr}
    \usepackage{xcolor}
    \definecolor{darkgray}{gray}{0.4} % Define a darker gray color
    \fancypagestyle{firstpage}{
        \fancyhf{} % Clear all headers and footers first
        \rhead{\textcolor{darkgray}{""" + month_year + r"""}}
        \renewcommand{\headrulewidth}{0pt} % No header rule
    }
    \thispagestyle{firstpage}
    """
    insertion_point = text_content.find(r'\begin{document}')
    if insertion_point != -1:
        text_content = text_content[:insertion_point] + header_settings + text_content[insertion_point:]
    return text_content


def set_professional_positions(text_content, doc):
    professional_positions_table_index = 3
    table_data = get_table_data(doc, professional_positions_table_index)

    positions = {
        "academic" : [
            {
                'Previous Employers with City/State\nIncluding U.S. Military\n(Most Recent First)': 'The Pennsylvania State University, University Park, PA', 
                'Work Performed: If Teacher, List Subjects Taught': 'Responsible for teaching undergraduate and graduate level classes in engineering and leading basic research in the area of computational biomechanics.',
                'Rank or Title': 'Professor of Mechanical and Biomedical Engineering (Courtesy)', 
                'Dates': '2024 - 2024'
            }
        ],
        "government" : [],
        "professional" : [
            {
                'Previous Employers with City/State\nIncluding U.S. Military\n(Most Recent First)': 'BrainSim Technologies Inc.', 
                'Work Performed: If Teacher, List Subjects Taught': 'Responsible for teaching undergraduate and graduate level classes in engineering and leading basic research in the area of computational biomechanics.',
                'Rank or Title': 'Founder and Chief Engineer', 
                'Dates': 'July 2019 - June 2024'
            }
        ]
    }

    for row in table_data:
        rank_or_title = row.get('Rank or Title', '')
        employer = row.get('Previous Employers with City/State\nIncluding U.S. Military\n(Most Recent First)', '')
        if "Associate Professor" in rank_or_title or "Assistant Professor" in rank_or_title:
            positions["academic"].append(row)
        elif "U.S. Army Research Laboratory" in employer:
            positions["government"].append(row)
        else:
            positions["professional"].append(row)

    latex_content = r"""
    \section*{PROFESSIONAL POSITIONS}
    """
    for category, entries in positions.items():
        latex_content += rf"""
    \subsection*{{{category.capitalize()}}}
    """
        for entry in entries:
            title_and_employer = entry.get('Rank or Title', '').replace(" (Courtesy)", "").strip() + ", " + entry.get('Previous Employers with City/State\nIncluding U.S. Military\n(Most Recent First)', '').strip()
            dates = entry.get('Dates', '').strip()
            years = re.findall(r'\b(\d{4})\b', dates)
            if years and len(years) == 2:
                current_year = str(datetime.datetime.now().year)
                if years[1] == current_year:
                    if 'associate' in entry.get('Rank or Title', '').lower():
                        formatted_date = f"{years[0]} - 2024"
                    else:
                        formatted_date = f"{years[0]} - Present"
                else:
                    formatted_date = f"{years[0]} - {years[1]}"
            elif years and len(years) == 1:
                current_year = str(datetime.datetime.now().year)
                if years[0] == current_year:
                    formatted_date = f"Present"
                else:
                    formatted_date = f"{years[0]}"
            else:
                formatted_date = ""

            if category == 'academic':
                latex_content += r"""
                \noindent """ + title_and_employer + " , " + r"\textbf{" + formatted_date + r"}" + r"""\vspace{0.25cm}
                """
            else:
                latex_content += r"""
                \noindent \parbox[t]{0.8\linewidth}{\raggedright """ + title_and_employer + r"""} \hfill \parbox[t]{0.2\linewidth}{\raggedleft """ + formatted_date + r"""} \\
                """
            
    text_content = text_content.replace("{{positions}}", latex_content)

    return text_content

def add_education(text_content, doc):
    education_table_index = 1
    table_data = get_table_data(doc, education_table_index)

    latex_content = r"""
    \section*{EDUCATION}
    """
    for row in table_data:
        result = row.get('Name and City/State of Institution', '').split(",")
        if len(result) >= 4:
            institute_name = f"{result[0]}, {result[1]}"
        else:
            institute_name = result[0]
        
        degree_and_date = row.get('Degrees - Dates', '').split(",")

        if len(degree_and_date) >= 2:
            degree = degree_and_date[0].strip()
            dates = degree_and_date[1].strip()
        else:
            degree = ''
            dates = ''

        major = row.get('Major Subjects', '').strip()

        years = re.findall(r'\b(\d{4})\b', dates)
        if years and len(years) == 2:
            current_year = str(datetime.datetime.now().year)
            if years[1] == current_year:
                formatted_date = f"{years[0]} - Present"
            else:
                formatted_date = f"{years[0]} - {years[1]}"
        elif years and len(years) == 1:
            current_year = str(datetime.datetime.now().year)
            if years[0] == current_year:
                formatted_date = f"Present"
            else:
                formatted_date = f"{years[0]}"
        else:
            formatted_date = ""

        if 'post-doctoral' in degree.lower():
            latex_content += r"""
            \noindent \parbox[t]{0.8\linewidth}{\raggedright """ + degree + ", " + institute_name + r"""} \hfill \parbox[t]{0.2\linewidth}{\raggedleft """ + formatted_date + r"""} \\
            \noindent \parbox[t]{0.8\linewidth}{\raggedright \textbf{Concentration:} """ + major + r"""} \\
            """
        else:
            latex_content += r"""
            \noindent \parbox[t]{0.8\linewidth}{\raggedright """ + degree + ", " + institute_name + r"""} \hfill \parbox[t]{0.2\linewidth}{\raggedleft """ + formatted_date + r"""} \\
            \noindent \parbox[t]{0.8\linewidth}{\raggedright \textbf{Major:} """ + major + r"""} \\
            """

    text_content = text_content.replace("{{education}}", latex_content)

    return text_content

def add_awards_and_honors(text_content, document_text):
    markers = [
        ("Honors or Awards for Scholarship, Research, or Creative Activities\n\nScholarship/Research",
         "Technology Transferred or Adapted in the Field"),
        ("Service, Professional\n\n",
         "EXTERNAL LETTERS OF ASSESSMENT"),
    ]

    relevant_text = ""
    for start_marker, end_marker in markers:
        relevant_text += extract_text_between_markers(document_text, start_marker, end_marker) + "\n\n"
        relevant_text.lstrip(',').strip()
    awards = [award.strip() for award in relevant_text.strip().split("\n\n") if award.strip()]

    latex_content = r"""
    \section*{AWARDS AND HONORS}
    """
    award_and_honors = []
    for award in awards:
        award_details = award.split("\n")[0].strip().split(".")
        if len(award_details) >= 2 :
            award_title = award_details[0].strip()
            award_date = award_details[1].strip()
        
        years = re.findall(r'\b(\d{4})\b', award_date)
        if years and len(years) >= 1:
            formated_award_date = years[0]
        
        award_and_honors.append({"award_title": award_title, "formated_award_date": formated_award_date})

    sorted_awards = sorted(award_and_honors, key=lambda x: x['formated_award_date'], reverse=True)
    for award in sorted_awards:
        latex_content += r"""
        \noindent \parbox[t]{0.8\linewidth}{\raggedright """ + award["award_title"] + r"""} \hfill \parbox[t]{0.2\linewidth}{\raggedleft """ + award["formated_award_date"] + r"""} \\
        """

    text_content = text_content.replace("{{awards_and_honor}}", latex_content)

    return text_content


def underline_students(text):
    student_pattern = r'([A-Z][a-zA-Z\s\.,]+) \((?:Primary Author|Author|Co-Author|Student Author)(?: -? (Graduate Student|Undergraduate Student|Postdoctoral Student))?\)'    
    underline_text = re.sub(student_pattern, r'\\underline{\1}', text)
    return underline_text

def format_publication_entry(publication):
    doi_match = re.search(r"DOI: (\S+)", publication)
    if doi_match:
        doi_url = f"Published. \\url{{https://doi.org/{doi_match.group(1)}}}"
        publication = re.sub(r"DOI: \S+", "", publication)
    else:
        doi_url = ""
    publication = re.sub(r'(Kraft,)( R\.\s*H\.)', r'\\textbf{\\textbf{\1}\2}', publication)
    

    formatted_entry = f"\\item {publication.strip()} {doi_url}".strip()
    formatted_entry = formatted_entry.replace('&', '&\n')
    formatted_entry = replace_special_characters(formatted_entry)
    formatted_entry = underline_students(formatted_entry)
    return formatted_entry

def replace_special_characters(text):
    replacements = {
        'á': r"\'a",
        'é': r"\'e",
        'í': r"\'i",
        'ó': r"\'o",
        'ú': r"\'u",
        'ñ': r"\~n",
        'ü': r"\"u",
        'Á': r"\'A",
        'É': r"\'E",
        'Í': r"\'I",
        'Ó': r"\'O",
        'Ú': r"\'U",
        'Ñ': r"\~N",
        'Ü': r"\"U",
        'ç': r"\c{c}",
        'Ç': r"\c{C}",
        'ö': r"\"o",
        'Ö': r"\"O",
    }
    
    for char, latex_equiv in replacements.items():
        text = text.replace(char, latex_equiv)
    
    return text

def extract_publications(text_content, document_text):
    start_marker_journal = "Journal Article"
    end_marker_journal = "Parts of Books"
    start_marker_conference = "Refereed Conference Proceedings"
    end_marker_conference = "Other Works"
    start_marker_preprint = "Pre-Print"
    end_marker_preprint = "Manuscripts Submitted for Publication"

    journal_publications = extract_text_between_markers(document_text, start_marker_journal, end_marker_journal)
    conference_publications = extract_text_between_markers(document_text, start_marker_conference, end_marker_conference)
    preprint_publications = extract_text_between_markers(document_text, start_marker_preprint, end_marker_preprint)

    latex_content = r"""
    \section*{PUBLICATIONS}
    """

    latex_content += r"\textit{Mentored student and postdoc co-authors are underlined.}"

    latex_content += r"""
    \subsection*{Journal Articles}
    \begin{enumerate}
    """

    for publication in journal_publications.split("\n\n"):
        if publication.strip():
            publication = format_publication_entry(publication)
            formatted_publication = publication.replace("&", r"\&")
            formatted_publication = re.sub(r'^\\item\s*\d+\.', r'\\item', formatted_publication)
            latex_content += f" {formatted_publication}\n"

    latex_content += r"""
    \end{enumerate}
    """

    latex_content += r"""
    \subsection*{Conference Proceedings}
    \begin{enumerate}
    """
    for publication in conference_publications.split("\n\n"):
        if publication.strip():
            doimatch = re.search(r"^(.*?)\(\w+-\d+\).*?(DOI: [\S]+)", publication)
            if doimatch:
                publication_part = doimatch.group(1)  # Text before the DOI part
                doi_part = doimatch.group(2)
                publication = publication_part + " " + doi_part
            
            publication = format_publication_entry(publication)
            formatted_publication = publication.replace("&", r"\&")

            latex_content += f"  {formatted_publication}\n"

    latex_content += r"""
    \end{enumerate}
    """

    latex_content += r"""
    \subsection*{Preprints and Technical Reports}
    \begin{enumerate}
    """
    for publication in preprint_publications.split("\n\n"):
        if publication.strip():
            if "Technical Report" == publication.strip():
                continue
            publication = format_publication_entry(publication)
            formatted_publication = publication.replace("&", r"\&")

            latex_content += f"  {formatted_publication}\n"

    latex_content += r"""
    \end{enumerate}
    """

    text_content = text_content.replace("{{publications}}", latex_content)

    return text_content

def extract_contract_project_and_grants(doc, text_content, document_text):
    start_marker_grants = 'Projects, Grants, Commissions, and Contracts'
    end_marker_grants = 'Pending'
    extracted_text = extract_text_between_markers(document_text, start_marker_grants, end_marker_grants)

    latex_content = r"""
    \section*{CONTRACT, FELLOWSHIPS, GRANTS AND SPONSORED RESEARCH}
    """

    award_table_index = 34
    for award_text in extracted_text.split("\n\n"):
        if award_text == 'Awarded':
            continue

        result = {item.split(':', 1)[0].strip(): item.split(':', 1)[1].strip() for item in award_text.split("\n")}
        table = doc.tables[award_table_index]
        table_data = {}

        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if ":" in text:
                    key, value = map(str.strip, text.split(":", 1))
                    table_data.update({key: value})
                else:
                    table_data.update({"Data": text})
            
        result.update(table_data)
        award_amount = re.search(r'\$\d{1,3}(,\d{1,3})*(\.\d{2})?', result.get('Award Amount',''))
        award_point = (
            f"{result.get('Principal Investigator','')} (Principal Investigator), "
            f"``{result.get('Project Title','')}''"
            f", Sponsored by {result.get('Agency','')}, {award_amount.group()}. "
            f"({result.get('Start Date', '')} - {result.get('End Date', '')})."
        )        
        award_point = award_point.replace("$", r"\$")
        award_point = award_point.replace("&", r"\&")
        award_point = award_point.replace("#", r"\#")

        latex_content += r"""
        \noindent """ + award_point + r"""\vspace{0.25cm}
        """
        award_table_index += 1

    text_content = text_content.replace("{{awarded}}", latex_content)

    return text_content

def extract_intellectual_property(text_content, document_text):
    start_marker = 'Patent Intellectual Property'
    end_marker = 'Impact in Society of Research Scholarship and Creative Accomplishment'
    extracted_text = extract_text_between_markers(document_text, start_marker, end_marker)

    latex_content = r"""
    \section*{INTELLECTUAL PROPERTY}
    """

    for intellectual_property in extracted_text.split("\n\n"):
        qoute_part = re.search(r'(.*?"[^"]*")', intellectual_property)
        if qoute_part:
            qoute_text = qoute_part.group(0).strip()
        date_part = re.search(r'\(application: (?:\d{4}|\w+\s\d{4})\)', intellectual_property)
        if date_part:
            date_text = date_part.group(0).strip()

        ip_text = f"{qoute_text}  {date_text}"

        latex_content += r"""
        \noindent """ + ip_text + r"""\vspace{0.25cm}
        """
    text_content = text_content.replace("{{Intellectual}}", latex_content)

    return text_content

def extract_directed_student_learning(text_content, document_text):
    phd_start_marker = 'Ph.D. Dissertation Advisor'
    phd_end_marker = 'Ph.D. Dissertation Committee Member'
    phd_dissertation_text = extract_text_between_markers(document_text, phd_start_marker, phd_end_marker)

    latex_content = r"""
    \section*{DIRECTED STUDENT LEARNING}
    \subsection*{Ph.D. Dissertation}
    \begin{enumerate}
    """

    for phd_dissertation in phd_dissertation_text.split("\n\n"):
        phd_match = re.search(r'^(.*)\n', phd_dissertation)
        if phd_match:
            phd_dissertation = phd_match.group(1).strip()

        phd_dissertation_parts = phd_dissertation.split("Ph.D.")
        author = phd_dissertation_parts[0].strip()
        phd_part = phd_dissertation_parts[1]
        phd_part_match = re.match(r"^(.*)\s\(([^)]+)\)\.$", phd_part)
        if phd_part_match:
            phd_text = phd_part_match.group(1).strip()
            phd_text = re.sub(r"Date Graduated:.*$", "", phd_text).strip()
            phd_date = phd_part_match.group(2).strip()
        
            phd_dissertation = f"{author} ``{phd_text}'', {phd_date}"
        
        latex_content += f"""
        \\item {phd_dissertation}
        """

    latex_content += r"""
    \end{enumerate}
    """

    master_start_marker = "Master's Thesis Advisor"
    master_end_marker = "Master's Thesis Committee Member"
    master_dissertation_text = extract_text_between_markers(document_text, master_start_marker, master_end_marker)

    latex_content += r"""
    \subsection*{Master’s Thesis}
    \begin{enumerate}
    """

    for master_dissertation in master_dissertation_text.split("\n\n"):
        master_dissertation_parts = master_dissertation.split("MS.")
        if len(master_dissertation_parts) >= 2:
            author = master_dissertation_parts[0].strip()
            master_part = master_dissertation_parts[1]
            master_match = re.match(r"^(.*)\s\(([^)]+)\)\.$", master_part)
            if master_match:
                master_text = master_match.group(1).strip()
                master_text = re.sub(r"Date Graduated:.*$", "", master_text).strip()
                master_date = master_match.group(2).strip()
            
                master_dissertation = f"{author} ``{master_text}'', {master_date}"

        latex_content += f"""
        \\item {master_dissertation}
        """

    latex_content += r"""
    \end{enumerate}
    """

    postdoctoral_start_marker = "Postdoctoral Mentorship Advisor"
    postdoctoral_end_marker = "Research Activity Advisor"
    postdoctoral_dissertation_text = extract_text_between_markers(document_text, postdoctoral_start_marker, postdoctoral_end_marker)

    latex_content += r"""
    \subsection*{Postdoctoral Mentorship}
    \begin{enumerate}
    """

    for postdoctoral in postdoctoral_dissertation_text.split("\n\n"):
        postdoctoral_match = re.match(r"^(.*)\s\(([^)]+)\)\.$", postdoctoral)
        if postdoctoral_match:
            postdoctoral_text = postdoctoral_match.group(1).strip()
            postdoctoral_date = postdoctoral_match.group(2).strip()
        
            postdoctoral = f"{postdoctoral_text}, {postdoctoral_date}"
        
        latex_content += f"""
        \\item {postdoctoral}
        """
    latex_content += r"""
    \end{enumerate}
    """
    undergraduate_start_marker = "Undergraduate Honors Thesis Advisor"
    undergraduate_end_marker = "THE SCHOLARSHIP OF Research and \nCreative Accomplishments"
    undergraduate_thesis_text = extract_text_between_markers(document_text, undergraduate_start_marker, undergraduate_end_marker)

    latex_content += r"""
    \subsection*{Undergraduate Honors Thesis}
    \begin{enumerate}
    """

    for undergraduate_thesis in undergraduate_thesis_text.split("\n\n"):
        undergraduate_match = re.search(r'^(.*)\n', undergraduate_thesis)
        if undergraduate_match:
            undergraduate_thesis = undergraduate_match.group(1).strip()
        
        undergraduate_thesis_parts = undergraduate_thesis.split("Undergraduate.")
        author = undergraduate_thesis_parts[0].strip()
        undergraduate_part = undergraduate_thesis_parts[1]
        undergraduate_part_match = re.match(r"^(.*)\s\(([^)]+)\)\.$", undergraduate_part)
        if undergraduate_part_match:
            undergraduate_text = undergraduate_part_match.group(1).strip()
            undergraduate_text = re.sub(r"Date Graduated:.*$", "", undergraduate_text).strip()
            undergraduate_text = re.sub(r"\.$", "", undergraduate_text)
            undergraduate_date = undergraduate_part_match.group(2).strip()
        
            undergraduate_thesis = f"{author} ``{undergraduate_text}'', {undergraduate_date}"
        
        latex_content += f"""
        \\item {undergraduate_thesis}
        """

    latex_content += r"""
    \end{enumerate}
    """

    text_content = text_content.replace("{{directedstudent}}", latex_content)

    return text_content

def create_points(college_data, start_text, latex_content):
    for college in college_data.split("\n\n"):
        college_match = re.search(r'^(.*)\n', college)
        if college_match:
            college = college_match.group(1).strip()
        college_parts = re.match(r"^(.*)\s\(([^)]+)\)\.$", college)
        if college_parts:
            college_text = college_parts.group(1).strip()
            college_date = college_parts.group(2).strip()
            result = college_text.split(",")
            qoute = result[0].replace("&", r"\&")
            role = result[-1].replace(".","").strip()
            college = f"{start_text}, {role}, ``{qoute}''. ({college_date})"

        latex_content += r"""
        \noindent """ + college + r"""\vspace{0.25cm}
        """
    return latex_content

def extract_services(text_content, document_text):
    committee_work_start_marker = 'Record of Committee Work at Department, Division, School, Campus, College, and University Levels'
    committee_work_end_marker = 'Record of Academic Leadership and Support Work (College Representative, Faculty Mentoring, Assessment Activities, etc.)'
    committee_work_text = extract_text_between_markers(document_text, committee_work_start_marker, committee_work_end_marker)

    academic_start_marker = 'Academic Leadership and Support Work'
    academic_work_end_marker = 'Participation in Development/Fundraising Activities'
    academic_work_text = extract_text_between_markers(document_text, academic_start_marker, academic_work_end_marker)

    participation_start_marker = 'Participation in Development/Fundraising Activities'
    participation_work_end_marker = 'Service to Society as a Representative of the University'
    participation_work_text = extract_text_between_markers(document_text, participation_start_marker, participation_work_end_marker)

    committee_work_colleges = extract_text_between_markers(committee_work_text, "College", "University")
    academic_colleges = extract_text_between_markers(academic_work_text, "\n\nCollege\n\n")
    competition_judging_colleges = extract_text_between_markers(participation_work_text, "\n\nCollege\n\n")

    committee_work_departments = extract_text_between_markers(committee_work_text, "Department", "College")

    committee_work_university = extract_text_between_markers(committee_work_text, "University")
    participation_university = extract_text_between_markers(participation_work_text, "University", "Competition Judging")

    professions = extract_text_between_markers(document_text, "Organizing Conferences and Service on Conference Committees", "Honors or Awards for Leadership and/or Service to the University, Community, or the Profession")

    societies = extract_text_between_markers(document_text, "Service to Governmental Agencies at the International, Federal, State, or Local Levels", "Service to the Disciplines and to the Profession")

    latex_content = r"""
    \section*{SERVICE}
    \subsection*{College}
    """
    latex_content = create_points(committee_work_colleges, "Committee Work", latex_content)
    latex_content = create_points(academic_colleges, "Academic Leadership and Support Work", latex_content)
    latex_content = create_points(competition_judging_colleges, "Competition Judging", latex_content)

    latex_content += r"""
    \subsection*{Department}
    """
    latex_content = create_points(committee_work_departments, "Committee Work", latex_content)
    latex_content += r"""
    \subsection*{University}
    """
    latex_content = create_points(committee_work_university, "Committee Work", latex_content)
    latex_content = create_points(participation_university, "Participation in Development/Fundraising Activities", latex_content)
    latex_content += r"""
    \subsection*{Profession}
    """
    for profession in professions.split("\n\n"):
        profession_match = re.search(r'^(.*)\n', profession)
        if profession_match:
            profession = profession_match.group(1).strip()
        profession_parts = re.match(r"^(.*?)(\(\w+ \d{4} - \w+ \d{4}\)\.)$", profession)
        if profession_parts:
            profession_text = profession_parts.group(1).strip()
            profession_date = profession_parts.group(2).strip()
            result = profession_text.split(",")
            if len(result) >= 2:
                profession_qoute = result[0].strip()
                institute = result[1].replace("&", r"\&").strip()
                profession = f"Organizing Conferences and Service on Conference Committees, ``{profession_qoute}'', {institute} ({profession_date})"
        
        latex_content += r"""
        \noindent """ + profession + r"""\vspace{0.25cm}
        """
    latex_content += r"""
    \subsection*{Society}
    """
    for society in societies.split("\n\n"):
        society_parts = re.match(r"^(.*)\s\(([^)]+)\)\.$", society)
        if society_parts:
            society_text = society_parts.group(1).strip()
            society_date = society_parts.group(2).strip()
            result = society_text.split(",")
            if len(result) >= 3:
                society_qoute = result[0].strip()
                society_institute = result[1].strip()
                society_role = result[2].replace(".", "").strip()
                society = f"Service to Governmental Agencies, {society_role}, ``{society_qoute}'', {society_institute}. ({society_date})"

        latex_content += r"""
        \noindent """ + society + r"""\vspace{0.25cm}
        """  
    text_content = text_content.replace("{{service}}", latex_content)

    return text_content

def extract_editorial_board_positions(text_content, document_text):
    editorial_board_start_marker = 'Outreach - Editorial Responsibilities'
    editorial_board_end_marker = 'Outreach - Peer Review of Grant Proposals'
    editorial_board_text = extract_text_between_markers(document_text, editorial_board_start_marker, editorial_board_end_marker)

    latex_content = r"""
    \section*{EDITORIAL BOARD POSITIONS}
    """
    for editorial_board in editorial_board_text.split("\n\n"):
        latex_content += r"""
        \noindent """ + editorial_board.strip() + r"""\vspace{0.25cm}
        """  
    text_content = text_content.replace("{{editorialboard}}", latex_content)

    return text_content

def extract_professional_memberships(text_content, document_text):
    professional_start_marker = 'Record of Membership in Professional and Learned Societies'
    professional_end_marker = 'Description of New Courses and/or Programs Developed, Including Service Learning and Outreach Courses'
    professional_text = extract_text_between_markers(document_text, professional_start_marker, professional_end_marker)

    latex_content = r"""
    \section*{PROFESSIONAL MEMBERSHIPS}
    """

    for professional in professional_text.split("\n\n"):
        if professional == "National":
            continue
        latex_content += r"""
        \noindent """ + professional.strip() + r"""\vspace{0.25cm}
        """  
    text_content = text_content.replace("{{professionalmembership}}", latex_content)

    return text_content

def main():
    
    filename = 'output.tex'
    doc_file = "CV_Data.docx"

    document_text = paragraphs_text(doc_file)

    doc =  Document(doc_file)

    create_template_latex_file(filename)

    with open(filename, 'r') as file:
        text2 = file.read()
        
    text2 = set_section_colors(text2)
      
    text2 = format_header(text2, doc)
    
    text2 = add_date_to_header(text2)

    text2 = set_professional_positions(text2, doc)

    text2 = add_education(text2, doc)

    text2 = add_awards_and_honors(text2, document_text)

    text2 = extract_publications(text2, document_text)

    text2 = extract_contract_project_and_grants(doc, text2, document_text)

    text2 = extract_intellectual_property(text2, document_text)

    text2 = extract_directed_student_learning(text2, document_text)

    text2 = extract_services(text2, document_text)

    text2 = extract_editorial_board_positions(text2, document_text)

    text2 = extract_professional_memberships(text2, document_text)

    with open("output.tex", 'w') as file:
        file.write(text2)

if __name__ == "__main__":
    main()